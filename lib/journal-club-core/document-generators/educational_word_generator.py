#!/usr/bin/env python3
"""
Readable formatter with clear visual separation and proper structure
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import re

def remove_emojis(text):
    """Remove all emojis/unicode icons from text"""
    if not text:
        return text
    # Remove emojis using regex (covers all emoji ranges)
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        u"\U0001F900-\U0001F9FF"  # supplemental symbols
        u"\U0001FA00-\U0001FA6F"  # extended symbols
        "]+", flags=re.UNICODE)
    return emoji_pattern.sub('', text)

def set_table_autofit(table):
    """Set table to auto-fit content - Simple approach"""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove any existing width settings
    for tblW in tblPr.findall(qn('w:tblW')):
        tblPr.remove(tblW)

    # Set table width to auto (type='auto')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)

    # Set table layout to autofit
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'autofit')

    # Remove column widths from all cells to allow auto-sizing
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.find(qn('w:tcPr'))
            if tcPr is not None:
                for tcW in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(tcW)

def add_shaded_box(doc, content, bg_color='F0F0F0'):
    """Add a shaded box for better visual separation"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    set_table_autofit(table)  # Auto-fit table to content
    cell = table.cell(0, 0)

    # Add content to cell
    p = cell.paragraphs[0]
    p.text = content
    p.paragraph_format.space_after = Pt(3)

    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    cell._element.get_or_add_tcPr().append(shading)

    # Add spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def parse_markdown_table(lines):
    """Parse markdown table into rows and columns"""
    table_data = []
    for line in lines:
        line = line.strip()
        if not line or '---' in line:
            continue
        if '|' not in line:
            continue
        # Split by | and clean - handle both | col | col | and col | col formats
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty strings
        if cells:
            table_data.append(cells)
    return table_data

def create_word_table_from_markdown(doc, table_data):
    """Create a styled Word table from parsed markdown table data"""
    print(f"[DEBUG] create_word_table_from_markdown called with {len(table_data) if table_data else 0} rows", file=sys.stderr)
    if not table_data or len(table_data) < 2:
        print(f"[DEBUG] Skipping table - insufficient data", file=sys.stderr)
        return

    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    set_table_autofit(table)

    # Fill header row (first row)
    header_cells = table.rows[0].cells
    for i, cell_text in enumerate(table_data[0]):
        cell = header_cells[i]
        cell.text = cell_text
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # Header background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '4472C4')
        cell._element.get_or_add_tcPr().append(shading)
        # White text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Fill data rows
    for row_idx in range(1, len(table_data)):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(table_data[row_idx]):
            cell = row_cells[col_idx]
            cell.text = cell_text
            # Format cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def is_ascii_table_separator(line):
    """DISABLED - Only markdown tables with | are supported now"""
    return False  # Disable ASCII table parsing to force markdown format

def parse_ascii_table(lines, start_idx):
    """Parse ASCII-style space-separated table
    Format:
      Column1   Column2   Column3
     -------- --------- ----------
      Value1   Value2    Value3
    """
    if start_idx >= len(lines) or start_idx + 1 >= len(lines):
        return None

    header_line = lines[start_idx].rstrip()
    sep_line = lines[start_idx + 1] if start_idx + 1 < len(lines) else ""

    if not is_ascii_table_separator(sep_line):
        return None

    print(f"[DEBUG parse_ascii_table] Header: {header_line[:80]}", file=sys.stderr)
    print(f"[DEBUG parse_ascii_table] Separator: {sep_line[:80]}", file=sys.stderr)

    # Find column positions from separator line
    sep_parts = sep_line.split()
    column_positions = []
    search_start = 0
    for part in sep_parts:
        pos = sep_line.find(part, search_start)
        column_positions.append((pos, pos + len(part)))
        search_start = pos + len(part)

    num_cols = len(column_positions)
    print(f"[DEBUG parse_ascii_table] Found {num_cols} columns at positions: {column_positions}", file=sys.stderr)

    # Extract headers based on column positions
    headers = []
    for i, (start, end) in enumerate(column_positions):
        # For last column, take rest of line
        if i == num_cols - 1:
            header_text = header_line[start:].strip()
        else:
            next_start = column_positions[i + 1][0]
            header_text = header_line[start:next_start].strip()
        headers.append(header_text if header_text else f"Col{i+1}")

    print(f"[DEBUG parse_ascii_table] Headers: {headers}", file=sys.stderr)

    # Extract data rows
    rows = []
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].rstrip()

        # Stop at empty line, separator, or next heading
        if not line or line.strip().startswith('#') or is_ascii_table_separator(line):
            break

        # Extract cell values based on column positions
        row_cells = []
        for col_idx, (start, end) in enumerate(column_positions):
            if col_idx == num_cols - 1:
                # Last column takes rest of line
                cell_text = line[start:].strip() if start < len(line) else ""
            else:
                next_start = column_positions[col_idx + 1][0]
                cell_text = line[start:next_start].strip() if start < len(line) else ""
            row_cells.append(cell_text)

        # Only add non-empty rows
        if any(cell.strip() for cell in row_cells):
            rows.append(row_cells)

        i += 1

    print(f"[DEBUG parse_ascii_table] Extracted {len(rows)} data rows", file=sys.stderr)

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'end_line': i
        }

    return None

def format_checklist_with_tables(doc, content):
    """Parse checklist content and render tables as actual Word tables"""

    # Use the centralized table detection function
    cleaned_content, detected_tables = detect_and_convert_tables(content)

    # Now process the content with tables replaced by placeholders
    lines = cleaned_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Handle headings
        if line.startswith('####'):
            heading_text = line.replace('####', '').strip()
            p = doc.add_paragraph(heading_text)
            run = p.runs[0] if p.runs else p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
        elif line.startswith('###'):
            heading_text = line.replace('###', '').strip()
            p = doc.add_paragraph(heading_text)
            run = p.runs[0] if p.runs else p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            # Add a horizontal rule
            doc.add_paragraph('_' * 80).runs[0].font.color.rgb = RGBColor(200, 200, 200)
            i += 1
        elif line.startswith('##'):
            heading_text = line.replace('##', '').strip()
            p = doc.add_paragraph(heading_text)
            run = p.runs[0] if p.runs else p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            i += 1
        # Handle table placeholders
        elif line.startswith('[TABLE_') and line.endswith(']'):
            table_num = int(line[7:-1])
            if table_num < len(detected_tables):
                table_data = detected_tables[table_num]
                all_rows = [table_data['headers']] + table_data['rows']
                create_word_table_from_markdown(doc, all_rows)
            i += 1
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80).runs[0].font.color.rgb = RGBColor(200, 200, 200)
            i += 1
        # Regular paragraph
        else:
            # Collect multi-line paragraph
            para_text = line
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '|', '---', '**')):
                para_text += ' ' + lines[i].strip()
                i += 1

            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)
                # Check for bold markers
                if para_text.startswith('**') and '**' in para_text[2:]:
                    p.runs[0].font.bold = True

def format_methods_section(doc, content):
    """Format methods section with table support"""

    print(f"[DEBUG format_methods_section] Processing Methods section", file=sys.stderr)

    # First, detect and extract tables
    cleaned_content, detected_tables = detect_and_convert_tables(content)
    print(f"[DEBUG format_methods_section] Found {len(detected_tables)} tables", file=sys.stderr)

    # Split into paragraphs
    paragraphs = cleaned_content.split('\n\n')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Check for table markers
        if para.startswith('[TABLE_') and para.endswith(']'):
            try:
                table_num = int(para[7:-1])
                if table_num < len(detected_tables):
                    create_proper_table(doc, detected_tables[table_num])
                    print(f"[DEBUG format_methods_section] Rendered table {table_num}", file=sys.stderr)
            except:
                pass

        # Check for headings (lines starting with ### or emoji)
        elif para.startswith('###') or any(emoji in para[:5] for emoji in ['📊', '🎲', '🙈', '📐']):
            # Remove ### and emoji, make it a heading
            heading_text = re.sub(r'^###\s*|^[📊🎲🙈📐]\s*', '', para)
            p = doc.add_paragraph(heading_text)
            run = p.runs[0] if p.runs else p.add_run(heading_text)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            # Regular paragraph
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0)

def process_qa_section(doc, lines):
    """Process a Q&A section with proper formatting"""

    text = '\n'.join(lines)

    # Look for Q&A patterns
    if 'What' in text or 'Why' in text or 'How' in text:
        # This looks like a question
        parts = text.split(':', 1)
        if len(parts) == 2:
            # Format as Q&A
            # Question in a box
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            set_table_autofit(table)  # Auto-fit table to content
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            run = p.add_run(parts[0] + ':')
            run.font.bold = True
            run.font.size = Pt(11)

            # Answer below with indentation
            p = cell.add_paragraph()
            p.add_run(parts[1].strip())
            p.paragraph_format.left_indent = Inches(0.3)

            # Light blue background for Q&A
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E6F3FF')
            cell._element.get_or_add_tcPr().append(shading)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        else:
            # Regular paragraph
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
    else:
        # Regular text - check for bullet points
        for line in lines:
            if line.strip():
                if line.strip().startswith('•') or line.strip().startswith('-'):
                    # Bullet point
                    p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(4)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Inches(0.25)

def add_separator_line(doc):
    """Add a visual separator line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    # Add bottom border to create a line effect
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._element.get_or_add_pPr().append(pBdr)

def restart_numbering(paragraph):
    """Restart numbering for a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        # Remove existing numPr
        pPr.remove(numPr)

    # Add new numPr with restart
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def create_readable_document(data, output_path):
    """Create a document with readable formatting"""

    # Create detailed generation log
    log_path = output_path.replace('.docx', '_generation_log.txt')
    def log(msg):
        with open(log_path, 'a', encoding='utf-8') as f:
            f.write(msg + '\n')

    log("="*80)
    log("DOCUMENT GENERATION LOG")
    log("="*80)

    # Save raw data for debugging
    with open(output_path.replace('.docx', '_debug.txt'), 'w', encoding='utf-8') as f:
        for section in data.get('sections', []):
            f.write(f"\n\n========== {section.get('heading', 'NO HEADING')} ==========\n")
            f.write(section.get('content', 'NO CONTENT')[:1000])

    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # NOTE: No title at top - Paper Basics section contains Title, Authors, Journal, DOI
    # The PICO clinical question should NOT appear as document title

    # Process sections
    section_count = 0
    log(f"\nProcessing {len(data.get('sections', []))} sections...")

    for section in data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')
        images = section.get('images', [])

        log(f"\n--- Section {section_count + 1}: {heading} ---")
        log(f"Content length: {len(content)} chars")
        log(f"Has images: {len(images)}")

        # Skip "Tables and Figures" section if no images and no meaningful content
        if 'Tables and Figures' in heading:
            if not images and (not content or 'No high-importance' in content or 'not automatically detected' in content):
                log("Skipping Tables and Figures section (no content)")
                continue

        # No page breaks - user wants continuous flow
        # if heading and section_count > 0:
        #     doc.add_page_break()

        # Add section heading with visual separation
        if heading:
            section_count += 1

            # Remove emojis from heading
            heading = remove_emojis(heading)

            # Add a separator before major sections
            if 'Critical Appraisal Checklist' in heading or 'Methods Deep-Dive' in heading:
                add_separator_line(doc)

            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)

        # Process content based on section type
        log(f"Processing content for: {heading}")

        # Remove emojis from all content
        if content:
            content = remove_emojis(content)

        if 'Critical Appraisal Checklist' in heading and content:
            log("✓ Matched Critical Appraisal Checklist - calling format_checklist_with_tables()")
            # Parse and render checklist tables as actual Word tables
            format_checklist_with_tables(doc, content)
            log("✓ format_checklist_with_tables() completed")

        elif 'Methods' in heading and content:
            log("✓ Matched Methods section - calling format_methods_section()")
            # Format methods section with Q&A structure
            format_methods_section(doc, content)
            log("✓ format_methods_section() completed")

        elif content:
            # Process regular content
            # Detect and convert tables
            cleaned_content, detected_tables = detect_and_convert_tables(content)

            # Check if this section contains Q&A format (multiple numbered questions)
            has_numbered_qa = bool(re.search(r'Q\s+A\s*\n\s*-+', cleaned_content))
            first_numbered_in_section = True

            # Process paragraphs
            paragraphs = cleaned_content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Check for table markers
                if para.startswith('[TABLE_') and para.endswith(']'):
                    try:
                        table_num = int(para[7:-1])
                        if table_num < len(detected_tables):
                            create_proper_table(doc, detected_tables[table_num])
                    except:
                        pass

                # Check if this paragraph contains multiple numbered lines (Q&A format)
                elif has_numbered_qa or (para.count('\n') > 2 and re.search(r'^\d+\s+.+\n', para, re.MULTILINE)):
                    # Split by single newlines and process each line
                    lines = para.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line or line in ['Q A', 'Q  A', '--- ---', '------']:
                            continue

                        # Match "1  Question text" or "1. Question text" or "1   Question text"
                        number_match = re.match(r'^(\d+)\s+(.+)', line)
                        if number_match:
                            text_without_number = number_match.group(2).strip()
                            p = doc.add_paragraph(text_without_number, style='List Number')
                            p.paragraph_format.space_after = Pt(4)

                            # Restart numbering on first item
                            if first_numbered_in_section:
                                restart_numbering(p)
                                first_numbered_in_section = False
                        else:
                            # Regular line
                            if line:
                                p = doc.add_paragraph(line)
                                p.paragraph_format.space_after = Pt(4)

                # Check for Result/Interpretation blocks
                elif 'Result:' in para or 'Interpretation:' in para:
                    # Create a visually distinct block
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)

                    if 'Result:' in para:
                        parts = para.split('Result:', 1)
                        if len(parts) > 1:
                            p = cell.paragraphs[0]
                            run = p.add_run('Result:')
                            run.font.bold = True
                            run.font.size = Pt(11)
                            p.add_run('\n' + parts[1].strip())

                            # Light gray background
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'E8E8E8')
                            cell._element.get_or_add_tcPr().append(shading)

                    if 'Interpretation:' in para:
                        parts = para.split('Interpretation:', 1)
                        if len(parts) > 1:
                            p = cell.add_paragraph() if 'Result:' in para else cell.paragraphs[0]
                            run = p.add_run('Interpretation:')
                            run.font.bold = True
                            run.font.italic = True
                            run.font.size = Pt(11)
                            p.add_run('\n' + parts[1].strip())

                    doc.add_paragraph().paragraph_format.space_after = Pt(8)

                else:
                    # Check if this is a single numbered list item
                    number_match = re.match(r'^(\d+)\.\s+(.+)', para)
                    if number_match:
                        # Extract text without number and use Word's auto-numbering
                        text_without_number = number_match.group(2)
                        p = doc.add_paragraph(text_without_number, style='List Number')
                        p.paragraph_format.space_after = Pt(8)

                        # Restart numbering on first item
                        if first_numbered_in_section:
                            restart_numbering(p)
                            first_numbered_in_section = False
                    else:
                        # Regular paragraph with proper spacing
                        p = doc.add_paragraph()
                        p.add_run(para)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(8)

                        # Add indentation for body text
                        if not para.startswith('•') and not para.startswith('-'):
                            p.paragraph_format.first_line_indent = Inches(0.25)

    # Save document
    doc.save(output_path)
    return output_path

def detect_and_convert_tables(content):
    """Detect table-like content (both markdown and text) and convert to structured data"""

    tables_found = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Check for markdown table separator (supports both |---|---| and ---|--- formats)
        if '---' in line and '|' in line:
            # Separator line found, check for header above
            if i > 0:
                header_line = lines[i-1].strip()
                if '|' in header_line:  # Changed from startswith('|')
                    print(f"[DEBUG detect_and_convert] Found table at line {i}, header: {header_line[:60]}", file=sys.stderr)
                    table_data = extract_markdown_table(lines, i-1)
                    if table_data:
                        tables_found.append(table_data)
                        content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")
                        print(f"[DEBUG detect_and_convert] Table extracted, total tables: {len(tables_found)}", file=sys.stderr)

        # Check for ASCII-style table separator (space-separated with dash separators)
        elif is_ascii_table_separator(line):
            if i > 0:
                header_line = lines[i-1].strip()
                if header_line and not header_line.startswith('#'):
                    print(f"[DEBUG detect_and_convert] Found ASCII table at line {i}, header: {header_line[:60]}", file=sys.stderr)
                    table_data = parse_ascii_table(lines, i-1)
                    if table_data:
                        # Convert to format expected by detect_and_convert
                        all_rows = [table_data['headers']] + table_data['rows']
                        # Build original text for replacement
                        original_lines = lines[i-1:table_data['end_line']]
                        original_text = '\n'.join(original_lines)
                        tables_found.append({
                            'headers': table_data['headers'],
                            'rows': table_data['rows'],
                            'original_text': original_text
                        })
                        content = content.replace(original_text, f"[TABLE_{len(tables_found)-1}]")
                        print(f"[DEBUG detect_and_convert] ASCII table extracted, total tables: {len(tables_found)}", file=sys.stderr)

        # Check for dash separators (old format, but NOT if it has | or is ASCII table)
        elif '-----' in line and '|' not in line and not is_ascii_table_separator(line):
            if i > 0:
                header_line = lines[i-1].strip()
                if header_line and not header_line.startswith('#'):
                    table_data = extract_table_from_position(lines, i-1)
                    if table_data:
                        tables_found.append(table_data)
                        content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")

        i += 1

    print(f"[DEBUG detect_and_convert] Total tables found: {len(tables_found)}", file=sys.stderr)
    return content, tables_found

def extract_markdown_table(lines, start_idx):
    """Extract a markdown-style table starting from header row
    Supports both | Col | Col | and Col | Col formats
    """

    headers = []
    rows = []
    original_text_lines = []

    # Get header line
    header_line = lines[start_idx].strip()
    if '|' not in header_line:
        return None

    original_text_lines.append(lines[start_idx])

    # Parse headers - handle both | Header | and Header | formats
    parts = header_line.split('|')
    # Remove empty strings from start/end
    parts = [p.strip() for p in parts if p.strip()]
    headers = parts

    if len(headers) < 2:
        return None

    print(f"[DEBUG extract_markdown_table] Parsed headers: {headers}", file=sys.stderr)

    # Skip separator line
    if start_idx + 1 < len(lines):
        original_text_lines.append(lines[start_idx + 1])

    # Get data rows
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if '|' not in line:  # Changed from startswith('|')
            break

        # Parse row - handle both formats
        parts = line.split('|')
        parts = [p.strip() for p in parts if p.strip()]
        row_data = parts

        if row_data and len(row_data) >= len(headers):
            rows.append(row_data[:len(headers)])
            original_text_lines.append(lines[i])

        i += 1

    print(f"[DEBUG extract_markdown_table] Extracted {len(rows)} rows", file=sys.stderr)

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def extract_table_from_position(lines, start_idx):
    """Extract a table starting from a given position"""

    headers = []
    rows = []
    original_text_lines = []

    # Get headers
    header_line = lines[start_idx].strip()
    original_text_lines.append(lines[start_idx])

    # Split headers
    headers = re.split(r'\s{2,}|\t', header_line)
    headers = [h.strip() for h in headers if h.strip()]

    if len(headers) < 2:
        return None

    # Skip separator line
    if start_idx + 1 < len(lines):
        original_text_lines.append(lines[start_idx + 1])

    # Get data rows
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('#'):
            break

        row_data = re.split(r'\s{2,}|\t', line)
        row_data = [d.strip() for d in row_data if d.strip()]

        if row_data:
            while len(row_data) < len(headers):
                row_data.append('')
            rows.append(row_data[:len(headers)])
            original_text_lines.append(lines[i])

        i += 1

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def create_proper_table(doc, table_data):
    """Create a real Word table with borders and autofit"""

    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])

    if not headers:
        return

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Properly set autofit to content
    set_table_autofit(table)

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # Gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D0D0D0')
        cell._element.get_or_add_tcPr().append(shading)

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data[:len(headers)]):
            cell = row.cells[i]
            cell.text = str(value) if value else ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Set cell width to auto
            tcPr = cell._element.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), '0')
            tcW.set(qn('w:type'), 'auto')

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

def main():
    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': 'Usage: script.py input.json output.docx'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        create_readable_document(data, output_path)
        print(json.dumps({'success': True}))

    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()