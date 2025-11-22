#!/usr/bin/env python3
"""
REAL Table Formatter - Actually converts ALL table-like content to proper Word tables
No more fake tables with dashes!
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import re

def detect_and_convert_tables(content):
    """Detect ALL table-like content and convert to structured data"""

    # Pattern 1: Tables with dashes (like "---- ---- ----")
    # Pattern 2: Tables with pipes (like "| col1 | col2 |")
    # Pattern 3: Text that looks like columnar data

    tables_found = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Check for dash separators
        if '-----' in line or '────' in line or '═══' in line:
            # This might be a table - look for headers above
            if i > 0:
                header_line = lines[i-1].strip()
                if header_line and not header_line.startswith('#'):
                    # Found a potential table
                    table_data = extract_table_from_position(lines, i-1)
                    if table_data:
                        tables_found.append(table_data)
                        # Replace the table text with a marker
                        content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")

        # Check for pipe-separated tables
        elif '|' in line and line.count('|') >= 2:
            table_data = extract_pipe_table(lines, i)
            if table_data:
                tables_found.append(table_data)
                content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")

        i += 1

    return content, tables_found

def extract_table_from_position(lines, start_idx):
    """Extract a table starting from a given position"""

    headers = []
    rows = []
    original_text_lines = []

    # Get headers
    header_line = lines[start_idx].strip()
    original_text_lines.append(lines[start_idx])

    # Split headers by multiple spaces or specific delimiters
    headers = re.split(r'\s{2,}|\t', header_line)
    headers = [h.strip() for h in headers if h.strip()]

    if len(headers) < 2:
        return None

    # Skip separator line
    if start_idx + 1 < len(lines):
        original_text_lines.append(lines[start_idx + 1])

    # Get data rows
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('#') or '═══' in line:
            break

        # Split row data
        row_data = re.split(r'\s{2,}|\t', line)
        row_data = [d.strip() for d in row_data if d.strip()]

        if row_data:
            # Ensure row has same number of columns as headers
            while len(row_data) < len(headers):
                row_data.append('')
            rows.append(row_data[:len(headers)])
            original_text_lines.append(lines[i])

        i += 1

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def extract_pipe_table(lines, start_idx):
    """Extract a pipe-separated table"""

    headers = []
    rows = []
    original_text_lines = []

    # Process lines that have pipes
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        original_text_lines.append(lines[i])

        # Remove leading and trailing pipes
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]

        # Split by pipes
        cells = [cell.strip() for cell in line.split('|')]

        # Skip separator rows (like |---|---|)
        if all(set(cell.replace('-', '').replace('=', '').strip()) == set() for cell in cells):
            i += 1
            continue

        if not headers:
            headers = cells
        else:
            rows.append(cells)

        i += 1

    if headers and rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def create_real_word_document(data, output_path):
    """Create a Word document with REAL tables and proper formatting"""

    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Process title
    title = data.get('title', 'Critical Appraisal')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)

    # Process sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')

        # Add heading
        if heading:
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # Process content - detect and convert ALL tables
        if content:
            cleaned_content, detected_tables = detect_and_convert_tables(content)

            # Process the content with table markers
            paragraphs = cleaned_content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Check if this is a table marker
                if para.startswith('[TABLE_') and para.endswith(']'):
                    # Extract table number
                    try:
                        table_num = int(para[7:-1])
                    except:
                        continue
                    if table_num < len(detected_tables):
                        create_proper_table(doc, detected_tables[table_num])

                # Check for Result/Interpretation pattern
                elif 'Result:' in para or 'Interpretation:' in para:
                    # Create styled result/interpretation blocks
                    if 'Result:' in para:
                        parts = para.split('Result:', 1)
                        if len(parts) > 1:
                            p = doc.add_paragraph()
                            run = p.add_run('Result: ')
                            run.font.bold = True
                            run.font.name = 'Calibri'
                            p.add_run(parts[1].strip())
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_after = Pt(6)
                            # Add light gray background
                            add_paragraph_shading(p, 'E8E8E8')

                    if 'Interpretation:' in para:
                        parts = para.split('Interpretation:', 1)
                        if len(parts) > 1:
                            p = doc.add_paragraph()
                            run = p.add_run('Interpretation: ')
                            run.font.bold = True
                            run.font.italic = True
                            run.font.name = 'Calibri'
                            p.add_run(parts[1].strip())
                            p.paragraph_format.space_after = Pt(8)

                else:
                    # Regular paragraph - NO JUSTIFIED ALIGNMENT!
                    p = doc.add_paragraph()
                    p.add_run(para)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # NOT JUSTIFIED!
                    p.paragraph_format.space_after = Pt(6)
                    # Only indent if it's body text, not lists
                    if para and not para.startswith('•') and not para.startswith('-') and (not para[0].isdigit() if para else True):
                        try:
                            p.paragraph_format.first_line_indent = Inches(0.25)
                        except:
                            pass  # Skip indentation if it fails

        # Handle checklist items as proper tables
        if 'checklist_items' in section and section['checklist_items']:
            table_data = {
                'headers': ['Item', 'Assessment', 'Rationale'],
                'rows': [[item.get('item', ''),
                         item.get('assessment', ''),
                         item.get('rationale', '')] for item in section['checklist_items']]
            }
            create_proper_table(doc, table_data)

    # Save document
    doc.save(output_path)
    return output_path

def create_proper_table(doc, table_data):
    """Create a REAL table with borders and proper formatting"""

    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])

    if not headers:
        return

    # Create table with Table Grid style (has all borders)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # Set column widths - skip if it causes issues
    try:
        available_width = Inches(6.5)
        col_width = available_width / len(headers) if len(headers) > 0 else Inches(2)
        for col in table.columns:
            col.width = col_width
    except:
        # If setting width fails, just use auto-fit
        pass

    # Add headers with gray background
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        # Add gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D0D0D0')
        cell._element.get_or_add_tcPr().append(shading)

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data[:len(headers)]):
            cell = row.cells[i]
            cell.text = str(value) if value else ''
            # Center align
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

def add_paragraph_shading(paragraph, color):
    """Add background shading to a paragraph"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading)

def main():
    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': 'Usage: script.py input.json output.docx'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        create_real_word_document(data, output_path)
        print(json.dumps({'success': True}))

    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()