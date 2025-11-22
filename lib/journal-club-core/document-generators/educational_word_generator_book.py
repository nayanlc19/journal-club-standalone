#!/usr/bin/env python3
"""
Book-style Educational Word Document Generator
Creates clean, readable documents with proper tables for all statistical data
"""

import json
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import io
from PIL import Image as PILImage

def sanitize_text(text, context='general'):
    """Sanitize text for Word document compatibility"""
    if not text:
        return ''

    # Handle unicode issues
    replacements = {
        '\u202f': ' ',  # Narrow no-break space (common in NEJM)
        '\u2011': '-',  # Non-breaking hyphen
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero-width space
        '\u2009': ' ',  # Thin space
        '\u200a': ' ',  # Hair space
        '\ufeff': '',   # Zero-width no-break space
        '\u2028': '\n', # Line separator
        '\u2029': '\n\n', # Paragraph separator
    }

    for char, replacement in replacements.items():
        text = text.replace(char, replacement)

    # Remove control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')

    return text.strip()


class BookStyleDocument:
    """Create a clean, book-style educational document"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.chapter_number = 0

    def setup_styles(self):
        """Setup clean, readable book styles"""
        styles = self.doc.styles

        # Body text style
        if 'BookBody' not in styles:
            body_style = styles.add_style('BookBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Georgia'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.line_spacing = 1.5
            body_style.paragraph_format.space_after = Pt(8)
            body_style.paragraph_format.first_line_indent = Inches(0.3)

        # Chapter title
        if 'ChapterTitle' not in styles:
            chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Arial'
            chapter_style.font.size = Pt(24)
            chapter_style.font.bold = True
            chapter_style.paragraph_format.space_before = Pt(36)
            chapter_style.paragraph_format.space_after = Pt(24)
            chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Section heading
        if 'SectionHeading' not in styles:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Arial'
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(12)

        # Subsection heading
        if 'SubsectionHeading' not in styles:
            sub_style = styles.add_style('SubsectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            sub_style.font.name = 'Arial'
            sub_style.font.size = Pt(13)
            sub_style.font.bold = True
            sub_style.paragraph_format.space_before = Pt(12)
            sub_style.paragraph_format.space_after = Pt(6)

    def add_chapter(self, title):
        """Add a new chapter with clean formatting"""
        if self.chapter_number > 0:
            self.doc.add_page_break()

        self.chapter_number += 1

        # Chapter heading
        chapter = self.doc.add_paragraph()
        chapter.style = 'ChapterTitle'
        chapter.add_run(f"Chapter {self.chapter_number}")

        # Chapter title
        title_para = self.doc.add_paragraph()
        title_para.style = 'ChapterTitle'
        title_para.add_run(title.upper())

        # Add separator line
        self.add_separator()

    def add_separator(self):
        """Add a subtle separator line"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("━" * 20)
        p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
        p.paragraph_format.space_after = Pt(18)

    def add_section(self, title):
        """Add a section heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'SectionHeading'
        heading.add_run(title)

    def add_subsection(self, title):
        """Add a subsection heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'SubsectionHeading'
        heading.add_run(title)

    def add_body_text(self, text):
        """Add body text with proper formatting"""
        if not text or not text.strip():
            return

        # Split into paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = self.doc.add_paragraph()
                p.style = 'BookBody'
                p.add_run(para_text)

    def parse_statistical_text(self, text):
        """Parse dense statistical text into structured data for tables"""
        tables = []

        # Pattern 1: Dose comparisons (e.g., "5 mg: -15.0% (95% CI -15.9% to -14.2%)")
        dose_pattern = r'(\d+\s*mg|placebo):\s*([-\d.]+\s*%)[^)]*\([^)]+\)'

        # Pattern 2: Outcome comparisons with vs/versus
        outcome_pattern = r'([^:]+):\s*([^v]+)\s+(?:vs\.?|versus)\s+([^,\n]+)'

        # Pattern 3: Statistical measures with values
        measure_pattern = r'([\w\s]+?):\s*([-\d.]+(?:\s*%)?)\s*(?:\(([^)]+)\))?'

        # Try to identify table type and parse accordingly
        lines = text.strip().split('\n')

        # Check if this looks like dose comparison data
        if any('mg:' in line or 'placebo:' in line.lower() for line in lines):
            table_data = self.parse_dose_response_table(text)
            if table_data:
                tables.append(table_data)

        # Check if this looks like outcome data
        elif any(indicator in text.lower() for indicator in ['primary', 'secondary', 'outcome', 'endpoint']):
            table_data = self.parse_outcome_table(text)
            if table_data:
                tables.append(table_data)

        # Check if this looks like baseline characteristics
        elif 'baseline' in text.lower() or 'characteristic' in text.lower():
            table_data = self.parse_baseline_table(text)
            if table_data:
                tables.append(table_data)

        return tables

    def parse_dose_response_table(self, text):
        """Parse dose-response data into table format"""
        rows = []
        lines = text.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Extract dose/treatment group
            dose_match = re.search(r'([\d.]+\s*mg|placebo|tirzepatide|finerenone|treatment|control)', line, re.I)
            if not dose_match:
                continue

            dose = dose_match.group(1)

            # Extract percentage/value
            value_match = re.search(r'([-\d.]+\s*%)', line)
            value = value_match.group(1) if value_match else ''

            # Extract confidence interval
            ci_match = re.search(r'\((95\s*%?\s*CI[^)]+)\)', line)
            ci = ci_match.group(1) if ci_match else ''

            # Extract p-value if present
            p_match = re.search(r'[pP]\s*[<=]\s*([\d.]+)', line)
            p_value = p_match.group(1) if p_match else ''

            if dose and (value or ci):
                rows.append([dose.title(), value, ci, p_value])

        if rows:
            return {
                'title': 'Treatment Effects by Dose',
                'headers': ['Dose/Group', 'Effect', '95% CI', 'P-value'],
                'rows': rows
            }
        return None

    def parse_outcome_table(self, text):
        """Parse outcome data into table format"""
        rows = []
        lines = text.strip().split('\n')

        current_outcome = None
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this is an outcome header
            if 'primary' in line.lower() or 'secondary' in line.lower():
                current_outcome = line
                continue

            # Parse the data line
            if ':' in line:
                parts = line.split(':', 1)
                measure = parts[0].strip()
                values = parts[1].strip() if len(parts) > 1 else ''

                # Try to split treatment vs control
                if 'vs' in values or 'versus' in values:
                    vs_parts = re.split(r'\s+(?:vs\.?|versus)\s+', values, maxsplit=1)
                    treatment = vs_parts[0].strip() if vs_parts else values
                    control = vs_parts[1].strip() if len(vs_parts) > 1 else ''
                    rows.append([measure, treatment, control])
                else:
                    rows.append([measure, values, ''])

        if rows:
            return {
                'title': 'Study Outcomes',
                'headers': ['Outcome Measure', 'Treatment', 'Control/Comparison'],
                'rows': rows
            }
        return None

    def parse_baseline_table(self, text):
        """Parse baseline characteristics into table format"""
        rows = []
        lines = text.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line or not ':' in line:
                continue

            parts = line.split(':', 1)
            characteristic = parts[0].strip()
            value = parts[1].strip() if len(parts) > 1 else ''

            # Clean up the characteristic name
            characteristic = characteristic.replace('•', '').replace('-', '').strip()

            if characteristic and value:
                rows.append([characteristic, value])

        if rows:
            return {
                'title': 'Baseline Characteristics',
                'headers': ['Characteristic', 'Value'],
                'rows': rows
            }
        return None

    def create_statistical_table(self, table_data):
        """Create a clean, readable Word table from parsed data"""
        if not table_data or not table_data.get('rows'):
            return

        # Add title
        if table_data.get('title'):
            title = self.doc.add_paragraph()
            title.add_run(table_data['title'])
            title.runs[0].font.bold = True
            title.runs[0].font.size = Pt(12)
            title.paragraph_format.space_before = Pt(12)
            title.paragraph_format.space_after = Pt(6)

        # Create table
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])

        # Ensure all rows have same number of columns
        num_cols = max(len(headers), max(len(row) for row in rows) if rows else 0)

        # Create the table
        table = self.doc.add_table(rows=len(rows) + (1 if headers else 0), cols=num_cols)
        table.style = 'Light Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add headers
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                if i < num_cols:
                    header_cells[i].text = header
                    # Bold headers
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                    # Shade header row
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), 'E8E8E8')
                    header_cells[i]._element.get_or_add_tcPr().append(shading)

        # Add data rows
        start_row = 1 if headers else 0
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[start_row + row_idx].cells
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < num_cols:
                    row_cells[col_idx].text = str(cell_value) if cell_value else ''
                    # Set font size for readability
                    for paragraph in row_cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        # Add spacing after table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    def add_statistical_section(self, content):
        """Process statistical content and create appropriate tables"""
        if not content:
            return

        # First check if we can parse this into tables
        tables = self.parse_statistical_text(content)

        if tables:
            # Create a table for each parsed dataset
            for table_data in tables:
                self.create_statistical_table(table_data)
        else:
            # If no tables could be parsed, try one more aggressive approach
            # Split by major sections and create simple two-column tables
            self.create_fallback_table(content)

    def create_fallback_table(self, content):
        """Create a simple two-column table as fallback for unparseable data"""
        lines = content.strip().split('\n')
        rows = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if ':' in line:
                parts = line.split(':', 1)
                rows.append([parts[0].strip(), parts[1].strip()])
            else:
                rows.append(['', line])

        if rows:
            table_data = {
                'title': 'Statistical Results',
                'headers': ['Measure', 'Value'],
                'rows': rows
            }
            self.create_statistical_table(table_data)

    def add_info_box(self, title, content, box_type='info'):
        """Add an information box with clean formatting"""
        # Create a single-cell table for the box
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Set box color based on type
        colors = {
            'info': 'E3F2FD',      # Light blue
            'warning': 'FFF3E0',    # Light orange
            'success': 'E8F5E9',    # Light green
            'pearl': 'F3E5F5',      # Light purple
            'key': 'FFF9C4'         # Light yellow
        }

        # Apply shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), colors.get(box_type, 'F5F5F5'))
        cell._element.get_or_add_tcPr().append(shading)

        # Add content
        p = cell.paragraphs[0]

        # Add icon and title
        icons = {
            'info': 'ℹ️',
            'warning': '⚠️',
            'success': '✅',
            'pearl': '💡',
            'key': '🔑'
        }

        icon = icons.get(box_type, '📌')
        p.add_run(f"{icon} {title}\n")
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(11)

        # Add content
        p.add_run(content)
        p.runs[1].font.size = Pt(10)

        # Add spacing
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)


def create_book_style_document(data):
    """Create a clean, book-style educational document"""
    book = BookStyleDocument()

    # Title page
    title_para = book.doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)

    title_run = title_para.add_run("JOURNAL CLUB")
    title_run.font.name = 'Arial Black'
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = book.doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Critical Appraisal & Educational Guide")
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(64, 64, 64)

    # Paper title
    paper_title = book.doc.add_paragraph()
    paper_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paper_title.paragraph_format.space_before = Pt(36)
    paper_title.add_run(sanitize_text(data.get('title', 'Untitled')))
    paper_title.runs[0].font.size = Pt(14)
    paper_title.runs[0].font.italic = True

    book.doc.add_page_break()

    # Table of Contents
    book.add_chapter("TABLE OF CONTENTS")

    toc_items = [
        "1. Executive Summary",
        "2. Study Design & Methods",
        "3. Statistical Analysis",
        "4. Primary Results",
        "5. Secondary Outcomes",
        "6. Critical Appraisal",
        "7. Clinical Implications",
        "8. Figures & Tables"
    ]

    for item in toc_items:
        p = book.doc.add_paragraph()
        p.style = 'BookBody'
        p.add_run(item)

    # Process each section
    sections = data.get('sections', [])

    for section in sections:
        heading = section.get('heading', '')
        content = sanitize_text(section.get('content', ''))

        if not heading or not content:
            continue

        # Start new chapter for major sections
        if any(key in heading.lower() for key in ['summary', 'method', 'result', 'discussion', 'appraisal']):
            book.add_chapter(heading)
        else:
            book.add_section(heading)

        # Check if this section contains statistical results
        if any(indicator in content for indicator in ['%', 'CI:', 'HR:', 'RR:', 'OR:', 'p=', 'P=', 'vs', 'versus']):
            # This content has statistics - parse and create tables
            book.add_statistical_section(content)
        else:
            # Regular content - check for special formatting
            if '•' in content or content.strip().startswith(('1.', '2.', '3.')):
                # This is a list
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        if line.startswith(('•', '-', '*')):
                            p = book.doc.add_paragraph(style='List Bullet')
                            p.add_run(line[1:].strip())
                        elif re.match(r'^\d+\.', line):
                            p = book.doc.add_paragraph(style='List Number')
                            p.add_run(re.sub(r'^\d+\.\s*', '', line))
                        else:
                            book.add_body_text(line)
            else:
                book.add_body_text(content)

        # Add clinical pearls if present
        if section.get('clinical_pearl'):
            book.add_info_box("Clinical Pearl", section['clinical_pearl'], 'pearl')

        # Add key points if present
        if section.get('key_points'):
            book.add_info_box("Key Points", section['key_points'], 'key')

    # Add figures section
    if data.get('figures'):
        book.add_chapter("FIGURES & TABLES")

        for idx, figure in enumerate(data.get('figures', []), 1):
            # Add figure title
            fig_title = book.doc.add_paragraph()
            fig_title.add_run(f"Figure {idx}. {sanitize_text(figure.get('title', 'Untitled'))}")
            fig_title.runs[0].font.bold = True
            fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add image if available
            if figure.get('imageBase64'):
                try:
                    img_data = base64.b64decode(figure['imageBase64'])
                    img = PILImage.open(io.BytesIO(img_data))

                    # Save to temp buffer
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Add to document
                    book.doc.add_picture(img_buffer, width=Inches(5))
                    last_paragraph = book.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Add caption
            if figure.get('caption'):
                caption = book.doc.add_paragraph()
                caption.add_run(sanitize_text(figure['caption']))
                caption.runs[0].font.size = Pt(9)
                caption.runs[0].font.italic = True
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            book.doc.add_paragraph()  # Spacing

    return book.doc


def main():
    """Main function to generate the document"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "error": "Usage: python educational_word_generator_book.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        # Load the input data
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Create the document
        doc = create_book_style_document(data)

        # Save the document
        doc.save(output_file)

        print(json.dumps({"success": True, "output": output_file}))

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()