#!/usr/bin/env python3
"""
PROPER Book Formatter - Actually creates tables and proper formatting
This time testing each component before integration
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import re

def add_shading_to_cell(cell, color):
    """Add background color to a cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def extract_tables_from_text(text):
    """Extract table data from text that looks like it should be a table"""
    tables = []

    # Pattern for Robson group data
    robson_pattern = r'Group\s+(\d+).*?(\d+\.?\d*)\s*%.*?CS.*?(\d+\.?\d*)\s*%.*?deliver'

    # Look for sections that contain table-like data
    if 'Robson' in text and '%' in text:
        # This is likely Robson classification data
        headers = ['Robson Group', 'Number of CS', '% of all CS', '% of all deliveries']
        rows = []

        # Find all groups mentioned
        lines = text.split('\n')
        for line in lines:
            if re.search(r'Group\s+\d+', line):
                # Extract group number and percentages
                match = re.search(r'Group\s+(\d+).*?(\d+).*?(\d+\.?\d*)\s*%.*?(\d+\.?\d*)\s*%', line)
                if match:
                    rows.append([f"Group {match.group(1)}", match.group(2),
                                f"{match.group(3)}%", f"{match.group(4)}%"])

        if rows:
            tables.append({'headers': headers, 'rows': rows, 'caption': 'Distribution by Robson Classification'})

    # Pattern for demographic data tables
    if 'Age' in text and ('mean' in text.lower() or 'years' in text):
        # Extract age statistics
        age_match = re.search(r'(?:mean|average).*?age.*?(\d+\.?\d*)\s*years', text, re.IGNORECASE)
        if age_match:
            headers = ['Variable', 'Value', 'Range']
            rows = [['Maternal Age (mean)', f"{age_match.group(1)} years", '']]

            # Look for age range
            range_match = re.search(r'(\d+)\s*-\s*(\d+)\s*years', text)
            if range_match:
                rows[0][2] = f"{range_match.group(1)}-{range_match.group(2)} years"

            tables.append({'headers': headers, 'rows': rows, 'caption': 'Demographic Characteristics'})

    return tables

def create_proper_document(data, output_path):
    """Create a properly formatted document with real tables and good spacing"""

    doc = Document()

    # Setup page
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)

    # Title
    title = data.get('title', 'Critical Appraisal')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Georgia'
    run.font.size = Pt(24)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(24)

    # Process sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')

        # CHAPTER heading
        if 'CHAPTER' in heading.upper():
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Arial Black'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(139, 0, 0)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)

        # Section heading
        elif heading:
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)

        # Process content - look for tables
        if content:
            # Check if content contains table data
            extracted_tables = extract_tables_from_text(content)

            if extracted_tables:
                # Add any text before the table
                lines = content.split('\n')
                intro_text = []
                for line in lines:
                    if not any(keyword in line for keyword in ['Group', '%', 'CS rate']):
                        intro_text.append(line)
                    else:
                        break

                if intro_text:
                    p = doc.add_paragraph()
                    p.add_run('\n'.join(intro_text))
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(8)

                # Add extracted tables
                for table_data in extracted_tables:
                    create_formatted_table(doc, table_data)

            else:
                # Check for Result/Interpretation pattern
                if 'Result:' in content and 'Interpretation:' in content:
                    # Split into result and interpretation
                    parts = content.split('Interpretation:')
                    result_part = parts[0].replace('Result:', '').strip()
                    interp_part = parts[1].strip() if len(parts) > 1 else ''

                    # Result box
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run('Result:')
                    run.font.bold = True
                    run.font.size = Pt(11)
                    p.add_run(f'\n{result_part}')
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(8)

                    # Add gray background
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E8E8E8')
                    p._element.get_or_add_pPr().append(shading)

                    # Interpretation
                    p = doc.add_paragraph()
                    run = p.add_run('Interpretation:')
                    run.font.bold = True
                    run.font.italic = True
                    run.font.size = Pt(11)
                    p.add_run(f'\n{interp_part}')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Inches(0.3)
                    p.paragraph_format.space_after = Pt(12)

                else:
                    # Regular paragraph
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            p = doc.add_paragraph()
                            p.add_run(para_text.strip())
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.space_after = Pt(8)
                            p.paragraph_format.line_spacing = 1.5
                            if not any(char in para_text[:50] for char in ['•', '-', '1.', '2.']):
                                p.paragraph_format.first_line_indent = Inches(0.3)

        # Handle checklist items as a table
        if 'checklist_items' in section and section['checklist_items']:
            items = section['checklist_items']
            table_data = {
                'headers': ['Criterion', 'Assessment', 'Explanation'],
                'rows': [[item.get('item', ''), item.get('assessment', ''),
                         item.get('rationale', '')] for item in items],
                'caption': f'{heading} Checklist'
            }
            create_formatted_table(doc, table_data)

    # Save document
    doc.save(output_path)
    # Don't print here - let main() handle output

def create_formatted_table(doc, table_data):
    """Create a properly formatted table with borders and shading"""

    headers = table_data['headers']
    rows = table_data['rows']
    caption = table_data.get('caption', '')

    # Add caption
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Table: {caption}')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    # Create table with borders
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'  # This ensures borders
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    table.allow_autofit = False
    total_width = Inches(6.5)
    col_width = total_width / len(headers)

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.width = col_width
        cell.text = str(header)

        # Format header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

        # Add gray background to header
        add_shading_to_cell(cell, 'D9D9D9')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data[:len(headers)]):
            cell = row.cells[i]
            cell.width = col_width
            cell.text = str(value) if value else ''

            # Center align
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Set font
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

def main():
    if len(sys.argv) < 3:
        # Test mode
        test_data = {
            'title': 'Critical Appraisal Test Document',
            'sections': [
                {
                    'heading': 'CHAPTER 8',
                    'content': ''
                },
                {
                    'heading': 'RESULTS INTERPRETATION (SIMPLE LANGUAGE)',
                    'content': '''The study looked at caesarean section rates across different groups.

Result: 165 CS out of 352 total deliveries → 46.9% (rounded to one decimal).
Interpretation: Almost one in two women who gave birth in this hospital during the study period had a caesarean section.'''
                },
                {
                    'heading': 'Distribution by Robson Classification',
                    'content': '''The Robson classification divides women into 10 groups based on pregnancy characteristics.

Group 1 had 10 CS representing 6.1% of all CS and 2.8% of deliveries.
Group 2 had 33 CS representing 20.0% of all CS and 9.4% of deliveries.
Group 5 had 41 CS representing 24.8% of all CS and 11.6% of deliveries.

The p-value < 0.001 indicates statistically significant differences between groups.'''
                }
            ]
        }

        output = 'test_proper_book.docx'
        create_proper_document(test_data, output)
        print(json.dumps({'success': True, 'message': f'Test document created: {output}'}))
    else:
        # Production mode
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            data = json.load(f)
        create_proper_document(data, sys.argv[2])
        print(json.dumps({'success': True}))

if __name__ == '__main__':
    main()