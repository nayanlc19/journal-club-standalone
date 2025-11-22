#!/usr/bin/env python3
"""
Clean Professional Educational Word Document Generator
Fixes alignment, justification, and formatting issues
"""

import json
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import io
from PIL import Image as PILImage

def sanitize_text(text, context='general'):
    """Sanitize text for Word document compatibility"""
    if not text:
        return ''

    # Handle unicode issues
    replacements = {
        '\u202f': ' ',  # Narrow no-break space (common in NEJM)
        '\u2011': '-',  # Non-breaking hyphen
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero-width space
        '\u2009': ' ',  # Thin space
        '\u200a': ' ',  # Hair space
        '\ufeff': '',   # Zero-width no-break space
        '\u2028': '\n', # Line separator
        '\u2029': '\n\n', # Paragraph separator
    }

    for char, replacement in replacements.items():
        text = text.replace(char, replacement)

    # Remove control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')

    return text.strip()


class CleanDocument:
    """Create a clean, professional educational document"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.chapter_number = 0

    def setup_styles(self):
        """Setup clean, professional styles WITHOUT justification"""
        styles = self.doc.styles

        # Clean body text - LEFT ALIGNED, NO JUSTIFICATION
        if 'CleanBody' not in styles:
            body_style = styles.add_style('CleanBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # NOT JUSTIFIED!
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.first_line_indent = Inches(0)  # NO INDENT!

        # Chapter title - clean and professional
        if 'ChapterTitle' not in styles:
            chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Calibri'
            chapter_style.font.size = Pt(20)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
            chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            chapter_style.paragraph_format.space_before = Pt(12)
            chapter_style.paragraph_format.space_after = Pt(18)

        # Section heading
        if 'SectionHeading' not in styles:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 0, 0)
            section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)

        # Subsection heading
        if 'SubHeading' not in styles:
            sub_style = styles.add_style('SubHeading', WD_STYLE_TYPE.PARAGRAPH)
            sub_style.font.name = 'Calibri'
            sub_style.font.size = Pt(12)
            sub_style.font.bold = True
            sub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_style.paragraph_format.space_before = Pt(6)
            sub_style.paragraph_format.space_after = Pt(3)

    def add_title_page(self, title):
        """Add a clean title page"""
        # Main title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(100)
        title_run = title_para.add_run("CRITICAL APPRAISAL")
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 139)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Pt(12)
        subtitle.add_run("Educational Analysis & Clinical Guide")
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(64, 64, 64)

        # Paper title
        paper_para = self.doc.add_paragraph()
        paper_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paper_para.paragraph_format.space_before = Pt(48)
        paper_para.add_run(sanitize_text(title))
        paper_para.runs[0].font.size = Pt(12)
        paper_para.runs[0].font.italic = True

        self.doc.add_page_break()

    def add_chapter(self, title):
        """Add a chapter with clean formatting"""
        if self.chapter_number > 0:
            self.doc.add_page_break()

        self.chapter_number += 1

        # Chapter heading
        chapter = self.doc.add_paragraph()
        chapter.style = 'ChapterTitle'
        chapter.add_run(f"{self.chapter_number}. {title}")

    def add_section(self, title):
        """Add a section heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'SectionHeading'
        heading.add_run(title)

    def add_subsection(self, title):
        """Add a subsection heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'SubHeading'
        heading.add_run(title)

    def add_text(self, text):
        """Add body text with CLEAN formatting (no justification!)"""
        if not text or not text.strip():
            return

        # Split into paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = self.doc.add_paragraph()
                p.style = 'CleanBody'
                p.add_run(para_text)

    def add_bullet_list(self, items):
        """Add a clean bullet list"""
        for item in items:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(item)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

    def parse_and_create_table(self, text):
        """Aggressively parse statistical text and create proper tables"""
        # Pattern matching for different types of data
        lines = text.strip().split('\n')
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for statistical patterns
            if any(indicator in line for indicator in ['%', 'CI', 'HR', 'RR', 'OR', 'p=', 'P=', 'vs', 'versus', '±', 'n=']):
                # This line has statistics - parse it
                if ':' in line:
                    parts = line.split(':', 1)
                    label = parts[0].strip().replace('•', '').replace('-', '').strip()
                    value = parts[1].strip() if len(parts) > 1 else ''

                    # Check if value has comparison
                    if 'vs' in value or 'versus' in value:
                        value_parts = re.split(r'\s+(?:vs\.?|versus)\s+', value)
                        if len(value_parts) >= 2:
                            table_data.append([label, value_parts[0].strip(), value_parts[1].strip()])
                        else:
                            table_data.append([label, value, ''])
                    else:
                        table_data.append([label, value, ''])
                else:
                    # No colon, might be a continuation or standalone value
                    table_data.append(['', line, ''])

        return table_data

    def create_clean_table(self, data, headers=None, title=None):
        """Create a clean, professional table"""
        if not data:
            return

        # Add title if provided
        if title:
            title_para = self.doc.add_paragraph()
            title_para.style = 'SubHeading'
            title_para.add_run(title)

        # Determine number of columns
        num_cols = max(len(row) for row in data) if data else len(headers) if headers else 2

        # Create table
        num_rows = len(data) + (1 if headers else 0)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light List'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Add headers
        row_idx = 0
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers[:num_cols]):
                header_cells[i].text = header
                # Format header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                # Light gray background for header
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F0F0F0')
                header_cells[i]._element.get_or_add_tcPr().append(shading)
            row_idx = 1

        # Add data rows
        for data_row in data:
            cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(data_row[:num_cols]):
                cells[col_idx].text = str(value) if value else ''
                # Format cells
                for paragraph in cells[col_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            row_idx += 1

        # Add spacing after table
        self.doc.add_paragraph()

    def add_info_box(self, title, content, color='blue'):
        """Add a clean info box"""
        # Create single-cell table for box
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Color schemes
        colors = {
            'blue': 'E8F4F8',
            'green': 'E8F5E9',
            'yellow': 'FFF9E6',
            'red': 'FFEBEE',
            'purple': 'F3E5F5'
        }

        # Apply shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), colors.get(color, 'F5F5F5'))
        cell._element.get_or_add_tcPr().append(shading)

        # Add content
        p = cell.paragraphs[0]
        # Title
        title_run = p.add_run(f"📌 {title}\n")
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        # Content
        content_run = p.add_run(content)
        content_run.font.size = Pt(10)

        # Add spacing
        self.doc.add_paragraph()

    def process_section_with_tables(self, content):
        """Process content and automatically create tables for statistical data"""
        if not content:
            return

        # Check if content has statistical data
        has_stats = any(indicator in content for indicator in ['%', 'CI:', 'HR:', 'RR:', 'OR:', 'p=', 'P=', 'vs', 'versus'])

        if has_stats:
            # Parse and create table
            table_data = self.parse_and_create_table(content)
            if table_data:
                # Determine headers based on content
                if any('vs' in str(row) or 'versus' in str(row) for row in table_data):
                    headers = ['Measure', 'Treatment/Result', 'Control/Comparison']
                else:
                    headers = ['Parameter', 'Value', 'Details']

                self.create_clean_table(table_data, headers=headers)
            else:
                # Fallback to text if parsing fails
                self.add_text(content)
        else:
            # Regular text content
            self.add_text(content)


def create_clean_document(data):
    """Create a clean, professional educational document"""
    doc = CleanDocument()

    # Title page
    doc.add_title_page(data.get('title', 'Untitled'))

    # Table of Contents
    doc.add_chapter("TABLE OF CONTENTS")

    toc_items = [
        "Executive Summary",
        "Study Design & Methods",
        "Statistical Analysis",
        "Primary Results",
        "Secondary Outcomes",
        "Critical Appraisal",
        "Clinical Implications"
    ]

    for i, item in enumerate(toc_items, 1):
        p = doc.doc.add_paragraph()
        p.style = 'CleanBody'
        p.add_run(f"{i}. {item}")

    # Process sections
    sections = data.get('sections', [])
    chapter_keywords = ['summary', 'introduction', 'method', 'result', 'discussion', 'appraisal', 'conclusion']

    for section in sections:
        heading = section.get('heading', '')
        content = sanitize_text(section.get('content', ''))

        if not heading:
            continue

        # Determine if this should be a chapter
        is_chapter = any(keyword in heading.lower() for keyword in chapter_keywords)

        if is_chapter:
            doc.add_chapter(heading.upper())
        else:
            doc.add_section(heading)

        # Process content based on type
        if content:
            # Check for lists
            if content.strip().startswith(('•', '-', '*', '1.', '2.')):
                lines = content.split('\n')
                items = []
                for line in lines:
                    line = line.strip()
                    if line:
                        # Remove bullet/number markers
                        line = re.sub(r'^[•\-\*]|\d+\.', '', line).strip()
                        if line:
                            items.append(line)
                if items:
                    doc.add_bullet_list(items)
            else:
                # Process for potential tables
                doc.process_section_with_tables(content)

        # Add clinical pearls if present
        if section.get('clinical_pearl'):
            doc.add_info_box("Clinical Pearl", section['clinical_pearl'], 'blue')

        # Add key points if present
        if section.get('key_points'):
            doc.add_info_box("Key Points", section['key_points'], 'green')

        # Add term definitions as table if present
        if section.get('term_definitions'):
            doc.add_section("Key Terms")
            term_data = [[term, definition] for term, definition in section['term_definitions'].items()]
            doc.create_clean_table(term_data, headers=['Term', 'Definition'])

    # Add figures if present
    if data.get('figures'):
        doc.add_chapter("FIGURES & TABLES")

        for idx, figure in enumerate(data.get('figures', []), 1):
            # Figure title
            fig_title = doc.doc.add_paragraph()
            fig_title.style = 'SubHeading'
            fig_title.add_run(f"Figure {idx}. {sanitize_text(figure.get('title', 'Untitled'))}")

            # Add image if available
            if figure.get('imageBase64'):
                try:
                    img_data = base64.b64decode(figure['imageBase64'])
                    img = PILImage.open(io.BytesIO(img_data))

                    # Save to buffer
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Add to document
                    doc.doc.add_picture(img_buffer, width=Inches(5))
                    last_paragraph = doc.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Add caption if present
            if figure.get('caption'):
                caption = doc.doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
                caption.add_run(sanitize_text(figure['caption']))
                caption.runs[0].font.size = Pt(10)
                caption.runs[0].font.italic = True

            doc.doc.add_paragraph()  # Spacing

    return doc.doc


def main():
    """Main function to generate the document"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "error": "Usage: python educational_word_generator_clean.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        # Load input data
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Create document
        doc = create_clean_document(data)

        # Save document
        doc.save(output_file)

        print(json.dumps({"success": True, "output": output_file}))

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()