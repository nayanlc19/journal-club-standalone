"""
Enhanced Educational Word Document Generator
Creates beautiful, notebook-style educational documents with excellent formatting
"""

import sys
import json
import base64
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

# Add the lib directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

try:
    from universal_text_sanitizer import sanitize_text, detect_publisher
except ImportError:
    def sanitize_text(text, context='general'):
        if not text:
            return text
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        text = text.replace('\u202f', ' ').replace('\u200b', '').replace('\u00a0', ' ')
        return text
    def detect_publisher(text):
        return 'unknown'


class NotebookStyleDocument:
    """Create a beautiful notebook-style educational document"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.setup_page_layout()
        self.toc_entries = []  # For table of contents

    def setup_page_layout(self):
        """Setup page margins and layout for notebook feel"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)  # Wider left margin for "notebook" feel
            section.right_margin = Cm(2)

    def setup_styles(self):
        """Setup custom styles for beautiful formatting"""
        styles = self.doc.styles

        # Main Title Style
        title_style = styles.add_style('MainTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Segoe UI'
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = RGBColor(25, 42, 86)  # Deep blue
        title_style.font.bold = True
        title_style.paragraph_format.space_after = Pt(24)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section Header Style (Large, colorful)
        section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Segoe UI Semibold'
        section_style.font.size = Pt(20)
        section_style.font.color.rgb = RGBColor(0, 102, 204)  # Bright blue
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(24)
        section_style.paragraph_format.space_after = Pt(12)
        section_style.paragraph_format.keep_with_next = True

        # Subsection Style
        subsection_style = styles.add_style('SubsectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        subsection_style.font.name = 'Segoe UI'
        subsection_style.font.size = Pt(16)
        subsection_style.font.color.rgb = RGBColor(51, 122, 183)  # Lighter blue
        subsection_style.font.bold = True
        subsection_style.paragraph_format.space_before = Pt(18)
        subsection_style.paragraph_format.space_after = Pt(8)

        # Key Concept Box Style
        key_concept_style = styles.add_style('KeyConcept', WD_STYLE_TYPE.PARAGRAPH)
        key_concept_style.font.name = 'Calibri'
        key_concept_style.font.size = Pt(11)
        key_concept_style.font.color.rgb = RGBColor(0, 0, 0)
        key_concept_style.paragraph_format.left_indent = Cm(0.5)
        key_concept_style.paragraph_format.right_indent = Cm(0.5)
        key_concept_style.paragraph_format.space_before = Pt(6)
        key_concept_style.paragraph_format.space_after = Pt(6)

        # Definition Style
        definition_style = styles.add_style('Definition', WD_STYLE_TYPE.PARAGRAPH)
        definition_style.font.name = 'Georgia'
        definition_style.font.size = Pt(11)
        definition_style.font.italic = True
        definition_style.font.color.rgb = RGBColor(102, 51, 153)  # Purple
        definition_style.paragraph_format.left_indent = Cm(1)

        # Note Style (margin notes)
        note_style = styles.add_style('NoteStyle', WD_STYLE_TYPE.PARAGRAPH)
        note_style.font.name = 'Calibri Light'
        note_style.font.size = Pt(10)
        note_style.font.color.rgb = RGBColor(105, 105, 105)  # Gray
        note_style.font.italic = True
        note_style.paragraph_format.left_indent = Cm(0.8)

        # Body Text Style
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.space_after = Pt(6)

    def add_colored_box(self, text, color_rgb=(230, 243, 255), border_color=(0, 102, 204)):
        """Add a colored box for important information"""
        # Create a table with one cell for the box effect
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Set cell background color
        cell = table.cell(0, 0)
        cell_xml = cell._element
        cell_props = cell_xml.get_or_add_tcPr()

        # Add shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '%02x%02x%02x' % color_rgb)
        cell_props.append(shading)

        # Add text to cell
        paragraph = cell.paragraphs[0]
        paragraph.text = text
        paragraph.style = 'KeyConcept'

        # Add spacing after box
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    def add_section_divider(self):
        """Add a visual section divider"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    def add_toc_entry(self, text, level=1):
        """Track entries for table of contents"""
        self.toc_entries.append({'text': text, 'level': level})

    def create_table_of_contents(self):
        """Create a table of contents at the beginning"""
        # Add TOC header
        toc_header = self.doc.add_paragraph()
        toc_header.style = 'SectionHeader'
        run = toc_header.add_run("📚 TABLE OF CONTENTS")
        run.font.color.rgb = RGBColor(25, 42, 86)

        # Add TOC entries
        for entry in self.toc_entries:
            p = self.doc.add_paragraph()
            if entry['level'] == 1:
                p.style = 'BodyText'
                p.add_run("▸ " + entry['text'])
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                p.style = 'NoteStyle'
                p.add_run("  ◦ " + entry['text'])
                p.paragraph_format.left_indent = Cm(1.5)

        # Add page break after TOC
        self.doc.add_page_break()

    def add_numbered_section(self, number, title, icon=""):
        """Add a numbered section with icon"""
        p = self.doc.add_paragraph()
        p.style = 'SectionHeader'

        # Add colored number box
        run_num = p.add_run(f" {number} ")
        run_num.font.color.rgb = RGBColor(255, 255, 255)
        run_num.font.bold = True

        # Add background color effect (using highlighting as workaround)
        from docx.enum.text import WD_COLOR_INDEX
        run_num.font.highlight_color = WD_COLOR_INDEX.BLUE

        # Add title with icon
        run_title = p.add_run(f"  {icon} {title}" if icon else f"  {title}")
        run_title.font.color.rgb = RGBColor(0, 102, 204)

        self.add_toc_entry(title, level=1)

    def add_subsection(self, title, icon=""):
        """Add a subsection with icon"""
        p = self.doc.add_paragraph()
        p.style = 'SubsectionHeader'
        run = p.add_run(f"{icon} {title}" if icon else title)
        self.add_toc_entry(title, level=2)

    def add_definition_box(self, term, definition):
        """Add a term definition in a styled box"""
        # Term
        p_term = self.doc.add_paragraph()
        p_term.style = 'SubsectionHeader'
        run = p_term.add_run(f"📖 {term}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(102, 51, 153)

        # Definition
        p_def = self.doc.add_paragraph()
        p_def.style = 'Definition'
        p_def.add_run(definition)

    def add_key_point(self, text):
        """Add a key point with emphasis"""
        p = self.doc.add_paragraph()
        # Add key point marker
        run_marker = p.add_run("💡 KEY POINT: ")
        run_marker.font.bold = True
        run_marker.font.color.rgb = RGBColor(255, 140, 0)  # Orange

        # Add text
        run_text = p.add_run(text)
        run_text.font.size = Pt(12)

        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def add_checklist_item(self, item, checked=False):
        """Add a checklist item"""
        p = self.doc.add_paragraph()
        checkbox = "☑" if checked else "☐"
        p.add_run(f"{checkbox} {item}")
        p.paragraph_format.left_indent = Cm(1)
        p.style = 'BodyText'

    def add_note(self, text):
        """Add a side note"""
        p = self.doc.add_paragraph()
        p.style = 'NoteStyle'
        run = p.add_run(f"📝 Note: {text}")


def create_enhanced_educational_document(data):
    """Create the enhanced educational document"""
    nb = NotebookStyleDocument()

    # Main Title
    title = nb.doc.add_paragraph()
    title.style = 'MainTitle'
    title.add_run(sanitize_text(data['title'], 'title'))

    # Subtitle
    subtitle = nb.doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Educational Critical Appraisal Guide")
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    subtitle.runs[0].font.italic = True

    nb.add_section_divider()

    # Process sections with enhanced formatting
    section_icons = {
        'Paper Basics': '📄',
        'Executive Summary': '📊',
        'Study Design': '🔬',
        'Historical Context': '📚',
        'Critical Appraisal Checklist': '✅',
        'Methods Deep-Dive': '🔍',
        'How This Study Was Done': '👥',
        'Results Interpretation': '📈',
        'Tables and Figures': '🖼️',
        'Criticisms': '⚠️',
        'Critical Analysis': '🎯',
        'Strengths': '💪',
        'Weaknesses': '❌',
        'Study Impact': '💥',
        'Clinical Implications': '🏥',
        'Defense Questions': '❓',
        'Funding': '💰',
        'Related Research': '🔗'
    }

    section_number = 1
    for section in data.get('sections', []):
        heading = sanitize_text(section.get('heading', ''))
        icon = section_icons.get(heading, '📌')

        # Add numbered section with icon
        nb.add_numbered_section(section_number, heading, icon)
        section_number += 1

        # Add content with proper formatting
        content = sanitize_text(section.get('content', ''))

        # Parse content for special formatting
        if content:
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check for special content types
                    if para.startswith('KEY:') or 'important' in para.lower()[:20]:
                        nb.add_key_point(para.replace('KEY:', '').strip())
                    elif para.startswith('NOTE:'):
                        nb.add_note(para.replace('NOTE:', '').strip())
                    elif ':' in para and len(para.split(':')[0]) < 50:
                        # Might be a definition
                        parts = para.split(':', 1)
                        if len(parts) == 2:
                            nb.add_definition_box(parts[0].strip(), parts[1].strip())
                        else:
                            p = nb.doc.add_paragraph(para)
                            p.style = 'BodyText'
                    else:
                        p = nb.doc.add_paragraph(para)
                        p.style = 'BodyText'

        # Add term definitions if present
        if section.get('term_definitions'):
            nb.add_subsection("Key Terms & Definitions", "📖")
            for term, definition in section['term_definitions'].items():
                nb.add_definition_box(sanitize_text(term), sanitize_text(definition))

        # Add teaching notes if present
        if section.get('teaching_notes'):
            nb.add_subsection("Teaching Notes", "👩‍🏫")
            nb.add_colored_box("Important points for understanding:")
            for note in section['teaching_notes']:
                nb.add_key_point(sanitize_text(note))

        # Add checklist items if present
        if section.get('checklist_items'):
            nb.add_subsection("Checklist", "☑️")
            for item in section['checklist_items']:
                assessment = item.get('assessment', 'Unclear')
                checked = assessment.lower() in ['present', 'yes', 'adequate']
                text = f"{item.get('item', '')} - {item.get('rationale', '')}"
                nb.add_checklist_item(sanitize_text(text), checked)

        # Add images if present
        if section.get('images'):
            nb.add_subsection("Visual Elements", "🖼️")
            for img in section['images']:
                # Add image title
                img_title = nb.doc.add_paragraph()
                img_title.style = 'SubsectionHeader'
                img_title.add_run(sanitize_text(img.get('title', 'Figure')))
                img_title.runs[0].font.size = Pt(14)

                # Try to add actual image
                if img.get('base64'):
                    try:
                        img_bytes = base64.b64decode(img['base64'])
                        img_stream = BytesIO(img_bytes)
                        nb.doc.add_picture(img_stream, width=Inches(5))
                    except:
                        nb.doc.add_paragraph("[Image could not be displayed]")

                # Add explanation
                if img.get('explanation'):
                    nb.add_note(sanitize_text(img['explanation']))

                nb.doc.add_paragraph()  # Spacing

        # Add section divider
        nb.add_section_divider()

    # Add metadata at the end
    if data.get('metadata'):
        nb.add_numbered_section(section_number, "Document Information", "ℹ️")
        meta = data['metadata']

        if meta.get('doi'):
            p = nb.doc.add_paragraph()
            p.add_run(f"DOI: {meta['doi']}")
            p.style = 'NoteStyle'

        if meta.get('journal'):
            p = nb.doc.add_paragraph()
            p.add_run(f"Journal: {meta['journal']}")
            p.style = 'NoteStyle'

        if meta.get('year'):
            p = nb.doc.add_paragraph()
            p.add_run(f"Year: {meta['year']}")
            p.style = 'NoteStyle'

    # Insert TOC at beginning (we have to do this last after collecting all entries)
    # This is a limitation of python-docx, so we'll skip actual TOC insertion for now
    # In production, you'd save and reload the document to insert at beginning

    return nb.doc


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python educational_word_generator_enhanced.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_json = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load input data
        with open(input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate enhanced document
        doc = create_enhanced_educational_document(data)

        # Save to file
        doc.save(output_path)

        print(json.dumps({
            "success": True,
            "message": f"Enhanced educational document created: {output_path}"
        }))

    except Exception as e:
        import traceback
        print(json.dumps({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }))
        sys.exit(1)


if __name__ == "__main__":
    main()