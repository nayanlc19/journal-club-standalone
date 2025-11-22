#!/usr/bin/env python3
"""
Final Educational Word Document Generator with Comprehensive Table Creation
Creates tables for ALL structured data: timelines, statistics, comparisons, etc.
"""

import json
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64
import io
from PIL import Image as PILImage

def sanitize_text(text, context='general'):
    """Sanitize text for Word document compatibility"""
    if not text:
        return ''

    # Handle unicode issues
    replacements = {
        '\u202f': ' ',  # Narrow no-break space (common in NEJM)
        '\u2011': '-',  # Non-breaking hyphen
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero-width space
        '\u2009': ' ',  # Thin space
        '\u200a': ' ',  # Hair space
        '\ufeff': '',   # Zero-width no-break space
        '\u2028': '\n', # Line separator
        '\u2029': '\n\n', # Paragraph separator
    }

    for char, replacement in replacements.items():
        text = text.replace(char, replacement)

    # Remove control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')

    return text.strip()


class FinalDocument:
    """Create a professional document with comprehensive table support"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.chapter_number = 0

    def setup_styles(self):
        """Setup clean professional styles"""
        styles = self.doc.styles

        # Clean body text - LEFT ALIGNED
        if 'Body' not in styles:
            body_style = styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_style.paragraph_format.line_spacing = 1.15
            body_style.paragraph_format.space_after = Pt(6)

        # Chapter title
        if 'Chapter' not in styles:
            chapter_style = styles.add_style('Chapter', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Calibri'
            chapter_style.font.size = Pt(18)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = RGBColor(0, 51, 102)
            chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            chapter_style.paragraph_format.space_before = Pt(12)
            chapter_style.paragraph_format.space_after = Pt(12)

        # Section heading
        if 'Heading1' not in styles:
            h1_style = styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = 'Calibri'
            h1_style.font.size = Pt(14)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)

        # Subsection heading
        if 'Heading2' not in styles:
            h2_style = styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = 'Calibri'
            h2_style.font.size = Pt(12)
            h2_style.font.bold = True
            h2_style.paragraph_format.space_before = Pt(6)
            h2_style.paragraph_format.space_after = Pt(3)

    def add_title_page(self, title):
        """Add title page"""
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(100)
        title_run = title_para.add_run("JOURNAL CLUB")
        title_run.font.name = 'Calibri Light'
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Critical Appraisal & Educational Analysis")
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(64, 64, 64)

        # Paper title
        paper_para = self.doc.add_paragraph()
        paper_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paper_para.paragraph_format.space_before = Pt(48)
        paper_para.add_run(sanitize_text(title))
        paper_para.runs[0].font.size = Pt(12)
        paper_para.runs[0].font.italic = True

        self.doc.add_page_break()

    def add_chapter(self, title):
        """Add chapter"""
        if self.chapter_number > 0:
            self.doc.add_page_break()

        self.chapter_number += 1
        chapter = self.doc.add_paragraph()
        chapter.style = 'Chapter'
        chapter.add_run(f"Chapter {self.chapter_number}: {title}")

    def add_section(self, title):
        """Add section heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'Heading1'
        heading.add_run(title)

    def add_subsection(self, title):
        """Add subsection heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'Heading2'
        heading.add_run(title)

    def add_text(self, text):
        """Add body text"""
        if not text or not text.strip():
            return

        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if para:
                p = self.doc.add_paragraph()
                p.style = 'Body'
                p.add_run(para)

    def parse_timeline_to_table(self, content):
        """Parse historical/timeline content into table format"""
        lines = content.strip().split('\n')
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Pattern 1: Year/Period followed by description
            year_match = re.match(r'^(\d{4}s?|\d{4}-\d{4}|\d{4})\s*[-:]?\s*(.+)', line)
            if year_match:
                period = year_match.group(1)
                description = year_match.group(2)
                table_data.append([period, description])
                continue

            # Pattern 2: Bullet points with dates
            bullet_match = re.match(r'^[•\-\*]\s*(\d{4}s?|\d{4}-\d{4})\s*[-:]?\s*(.+)', line)
            if bullet_match:
                period = bullet_match.group(1)
                description = bullet_match.group(2)
                table_data.append([period, description])
                continue

            # Pattern 3: Any line with a year in it
            if re.search(r'\b(19|20)\d{2}\b', line):
                # Try to split on common delimiters
                if ':' in line:
                    parts = line.split(':', 1)
                    table_data.append([parts[0].strip(), parts[1].strip()])
                elif '-' in line and line.index('-') < 20:
                    parts = line.split('-', 1)
                    table_data.append([parts[0].strip(), parts[1].strip()])
                else:
                    table_data.append(['', line])

        return table_data if table_data else None

    def parse_statistics_to_table(self, content):
        """Parse statistical content into table format - AGGRESSIVE VERSION"""
        lines = content.strip().split('\n')
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Remove bullet points
            line = re.sub(r'^[•\-\*]\s*', '', line)

            # Skip lines that are just headers or very short
            if len(line) < 10 or line.endswith(':'):
                continue

            # AGGRESSIVE: Any line with percentages or statistical indicators
            if any(indicator in line for indicator in ['%', 'CI', 'HR', 'RR', 'OR', 'p=', 'P=', 'p <', 'p<', '±', '95']):
                # Pattern 1: Measure: value vs value
                vs_match = re.search(r'(.+?):\s*([^v]+)\s+(?:vs\.?|versus)\s+(.+)', line, re.IGNORECASE)
                if vs_match:
                    measure = vs_match.group(1).strip()
                    treatment = vs_match.group(2).strip()
                    control = vs_match.group(3).strip()
                    table_data.append([measure, treatment, control])
                    continue

                # Pattern 2: Value with CI in parentheses anywhere
                if '(' in line and ')' in line:
                    # Extract the main part and the CI part
                    parts = re.split(r'\s*\([^)]+\)', line, 1)
                    if parts:
                        main_part = parts[0].strip()
                        ci_part = re.search(r'\(([^)]+)\)', line)
                        if ci_part:
                            ci_text = ci_part.group(1).strip()
                            # Check if this has a label
                            if ':' in main_part:
                                label_parts = main_part.split(':', 1)
                                table_data.append([label_parts[0].strip(), label_parts[1].strip(), ci_text])
                            else:
                                table_data.append(['', main_part, ci_text])
                            continue

                # Pattern 3: Colon-separated values
                if ':' in line:
                    parts = line.split(':', 1)
                    label = parts[0].strip()
                    value = parts[1].strip()
                    # Extract p-value if present
                    p_val = ''
                    p_val_match = re.search(r'p\s*[<=]\s*([\d.]+)', value, re.IGNORECASE)
                    if p_val_match:
                        p_val = f"p = {p_val_match.group(1)}"
                        value = re.sub(r'[,;]\s*p\s*[<=]\s*[\d.]+', '', value).strip()
                    table_data.append([label, value, p_val])
                    continue

                # Pattern 4: Just add the line if it has statistics
                table_data.append(['', line, ''])

        # If we have data, return it
        return table_data if len(table_data) > 0 else None

    def parse_baseline_to_table(self, content):
        """Parse baseline characteristics into table format"""
        lines = content.strip().split('\n')
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Remove bullet points
            line = re.sub(r'^[•\-\*]\s*', '', line)

            # Split on colon if present
            if ':' in line:
                parts = line.split(':', 1)
                characteristic = parts[0].strip()
                value = parts[1].strip()
                table_data.append([characteristic, value])
            elif '=' in line:
                parts = line.split('=', 1)
                characteristic = parts[0].strip()
                value = parts[1].strip()
                table_data.append([characteristic, value])
            else:
                table_data.append(['', line])

        return table_data if table_data else None

    def parse_outcomes_to_table(self, content):
        """Parse outcomes/endpoints into table format"""
        lines = content.strip().split('\n')
        table_data = []
        current_category = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this is a category header
            if line.lower().startswith(('primary', 'secondary', 'tertiary', 'exploratory')):
                current_category = line
                continue

            # Remove numbering and bullets
            line = re.sub(r'^[\d.]+\s*', '', line)
            line = re.sub(r'^[•\-\*]\s*', '', line)

            # Add to table with category if applicable
            if current_category:
                table_data.append([current_category, line])
                current_category = None  # Only use once
            else:
                table_data.append(['', line])

        return table_data if table_data else None

    def create_table(self, data, headers=None, title=None, style='Light Grid'):
        """Create a professional table"""
        if not data:
            return

        # Add title if provided
        if title:
            title_para = self.doc.add_paragraph()
            title_para.style = 'Heading2'
            title_para.add_run(title)

        # Create table
        num_cols = max(len(row) for row in data) if data else 3
        num_rows = len(data) + (1 if headers else 0)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Add headers
        row_idx = 0
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers[:num_cols]):
                header_cells[i].text = header
                # Bold headers
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                # Background color
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'E8E8E8')
                header_cells[i]._element.get_or_add_tcPr().append(shading)
            row_idx = 1

        # Add data
        for data_row in data:
            cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(data_row[:num_cols]):
                cells[col_idx].text = str(value) if value else ''
                # Format cells
                for para in cells[col_idx].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.size = Pt(10)
            row_idx += 1

        # Add spacing after table
        self.doc.add_paragraph()

    def process_content_intelligently(self, content, section_title=''):
        """Intelligently process content and create tables where appropriate"""
        if not content:
            return

        content = content.strip()
        section_lower = section_title.lower()

        # Debug logging
        print(f"[TABLE DEBUG] Processing section: {section_title}")
        has_stats = any(indicator in content for indicator in ['%', 'CI', 'HR', 'RR', 'OR', 'p=', 'P='])
        if has_stats:
            print(f"[TABLE DEBUG] Section has statistics indicators")

        # Timeline/Historical context
        if any(word in section_lower for word in ['timeline', 'history', 'historical', 'evolution', 'development']):
            table_data = self.parse_timeline_to_table(content)
            if table_data:
                print(f"[TABLE DEBUG] Creating timeline table with {len(table_data)} rows")
                self.create_table(
                    table_data,
                    headers=['Year/Period', 'Event/Development'],
                    title='Timeline'
                )
                return

        # Statistical results - BE MORE AGGRESSIVE
        # Don't just check section title, also check content
        if has_stats or any(word in section_lower for word in ['result', 'outcome', 'finding', 'efficacy', 'statistical', 'primary', 'secondary']):
            table_data = self.parse_statistics_to_table(content)
            if table_data:
                print(f"[TABLE DEBUG] Creating statistics table with {len(table_data)} rows")
                # Determine appropriate headers
                if any('vs' in str(row) or 'versus' in str(row) for row in table_data):
                    headers = ['Measure', 'Treatment', 'Control/Comparison']
                else:
                    headers = ['Parameter', 'Value', 'CI/P-value']

                self.create_table(
                    table_data,
                    headers=headers,
                    title='Statistical Results'
                )
                return
            else:
                print(f"[TABLE DEBUG] No table data parsed from statistics")

        # Baseline characteristics
        if any(word in section_lower for word in ['baseline', 'characteristic', 'demographic', 'population']):
            table_data = self.parse_baseline_to_table(content)
            if table_data:
                self.create_table(
                    table_data,
                    headers=['Characteristic', 'Value'],
                    title='Baseline Characteristics'
                )
                return

        # Outcomes/Endpoints
        if any(word in section_lower for word in ['endpoint', 'primary', 'secondary']):
            table_data = self.parse_outcomes_to_table(content)
            if table_data:
                self.create_table(
                    table_data,
                    headers=['Type', 'Outcome'],
                    title='Study Endpoints'
                )
                return

        # Methods with structured data
        if 'method' in section_lower and ':' in content:
            lines = content.split('\n')
            if len([l for l in lines if ':' in l]) > 2:
                table_data = self.parse_baseline_to_table(content)
                if table_data:
                    self.create_table(
                        table_data,
                        headers=['Parameter', 'Description'],
                        title='Study Design'
                    )
                    return

        # PICO framework
        if any(word in content.lower() for word in ['population:', 'intervention:', 'comparator:', 'outcome:']):
            table_data = self.parse_baseline_to_table(content)
            if table_data:
                self.create_table(
                    table_data,
                    headers=['PICO Element', 'Description'],
                    title='PICO Framework'
                )
                return

        # Risk of Bias
        if 'bias' in section_lower or 'rob' in section_lower:
            if ':' in content:
                table_data = self.parse_baseline_to_table(content)
                if table_data:
                    self.create_table(
                        table_data,
                        headers=['Domain', 'Assessment'],
                        title='Risk of Bias Assessment'
                    )
                    return

        # Lists (if has multiple bullet points or numbers)
        if content.count('•') > 2 or content.count('-') > 2 or re.search(r'\n\d+\.', content):
            lines = content.split('\n')
            list_items = []
            for line in lines:
                line = line.strip()
                if line:
                    # Remove bullets/numbers
                    line = re.sub(r'^[•\-\*]|\d+\.', '', line).strip()
                    if line:
                        list_items.append(line)

            # If items have consistent structure with colons, make a table
            if list_items and len([l for l in list_items if ':' in l]) > len(list_items) / 2:
                table_data = []
                for item in list_items:
                    if ':' in item:
                        parts = item.split(':', 1)
                        table_data.append([parts[0].strip(), parts[1].strip()])
                    else:
                        table_data.append(['', item])

                if table_data:
                    self.create_table(
                        table_data,
                        headers=['Item', 'Description']
                    )
                    return
            else:
                # Create as bullet list
                for item in list_items:
                    p = self.doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
                return

        # Default: add as regular text
        self.add_text(content)

    def add_info_box(self, title, content, color='blue'):
        """Add an info box"""
        # Create single-cell table
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Color schemes
        colors = {
            'blue': 'E3F2FD',
            'green': 'E8F5E9',
            'yellow': 'FFF9C4',
            'red': 'FFEBEE',
            'gray': 'F5F5F5'
        }

        # Apply shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), colors.get(color, 'F5F5F5'))
        cell._element.get_or_add_tcPr().append(shading)

        # Add content
        p = cell.paragraphs[0]
        title_run = p.add_run(f"💡 {title}\n")
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        content_run = p.add_run(content)
        content_run.font.size = Pt(10)

        self.doc.add_paragraph()


def create_final_document(data):
    """Create the final professional document with comprehensive tables"""
    doc = FinalDocument()

    # Title page
    doc.add_title_page(data.get('title', 'Untitled'))

    # Table of Contents
    doc.add_chapter("TABLE OF CONTENTS")

    toc_items = [
        "1. Executive Summary",
        "2. Historical Context & Timeline",
        "3. Study Design (PICO)",
        "4. Methods & Statistical Analysis",
        "5. Primary Results",
        "6. Secondary Outcomes",
        "7. Critical Appraisal",
        "8. Clinical Implications",
        "9. Figures & Tables"
    ]

    for item in toc_items:
        p = doc.doc.add_paragraph()
        p.style = 'Body'
        p.add_run(item)

    # Process sections
    sections = data.get('sections', [])

    for section in sections:
        heading = section.get('heading', '')
        content = sanitize_text(section.get('content', ''))

        if not heading:
            continue

        # Determine if this should be a chapter
        chapter_keywords = ['summary', 'introduction', 'background', 'method', 'result',
                          'discussion', 'conclusion', 'appraisal', 'implication']

        is_chapter = any(keyword in heading.lower() for keyword in chapter_keywords)

        if is_chapter:
            doc.add_chapter(heading.upper())
        else:
            doc.add_section(heading)

        # Process content intelligently
        if content:
            doc.process_content_intelligently(content, heading)

        # Add clinical pearls
        if section.get('clinical_pearl'):
            doc.add_info_box("Clinical Pearl", section['clinical_pearl'], 'blue')

        # Add key points
        if section.get('key_points'):
            doc.add_info_box("Key Points", section['key_points'], 'green')

        # Add term definitions as table
        if section.get('term_definitions'):
            doc.add_section("Key Terminology")
            term_data = [[term, definition] for term, definition in section['term_definitions'].items()]
            doc.create_table(
                term_data,
                headers=['Term', 'Definition'],
                title='Glossary'
            )

    # Add figures section
    if data.get('figures'):
        doc.add_chapter("FIGURES & TABLES")

        for idx, figure in enumerate(data.get('figures', []), 1):
            # Figure title
            fig_title = doc.doc.add_paragraph()
            fig_title.style = 'Heading2'
            fig_title.add_run(f"Figure {idx}. {sanitize_text(figure.get('title', 'Untitled'))}")

            # Add image if available
            if figure.get('imageBase64'):
                try:
                    img_data = base64.b64decode(figure['imageBase64'])
                    img = PILImage.open(io.BytesIO(img_data))

                    # Save to buffer
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Add to document
                    doc.doc.add_picture(img_buffer, width=Inches(5.5))
                    last_paragraph = doc.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image: {e}")

            # Add caption
            if figure.get('caption'):
                caption = doc.doc.add_paragraph()
                caption.add_run(sanitize_text(figure['caption']))
                caption.runs[0].font.size = Pt(10)
                caption.runs[0].font.italic = True
                caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

            doc.doc.add_paragraph()

    return doc.doc


def main():
    """Main function"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "error": "Usage: python educational_word_generator_final.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        # Load data
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Create document
        doc = create_final_document(data)

        # Save
        doc.save(output_file)

        print(json.dumps({"success": True, "output": output_file}))

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()