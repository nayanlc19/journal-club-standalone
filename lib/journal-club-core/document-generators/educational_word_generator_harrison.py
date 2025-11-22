"""
Harrison's Textbook Style Educational Document Generator
Creates professional medical textbook-style documents with proper formatting
Inspired by Harrison's Principles of Internal Medicine
"""

import sys
import json
import base64
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Add the lib directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

try:
    from universal_text_sanitizer import sanitize_text, detect_publisher
except ImportError:
    def sanitize_text(text, context='general'):
        if not text:
            return text
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        text = text.replace('\u202f', ' ').replace('\u200b', '').replace('\u00a0', ' ')
        return text
    def detect_publisher(text):
        return 'unknown'


class HarrisonStyleDocument:
    """Create a Harrison's Textbook of Medicine style document"""

    def __init__(self):
        self.doc = Document()
        self.setup_page_layout()
        self.setup_styles()
        self.chapter_number = 0

    def setup_page_layout(self):
        """Setup page layout for medical textbook style"""
        sections = self.doc.sections
        for section in sections:
            # Standard textbook margins
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            # Page size (Letter)
            section.page_height = Cm(27.94)
            section.page_width = Cm(21.59)

            # Headers and footers space
            section.header_distance = Cm(1.27)
            section.footer_distance = Cm(1.27)

    def setup_styles(self):
        """Setup Harrison's-style formatting"""
        styles = self.doc.styles

        # Chapter Number Style
        try:
            chapter_num_style = styles.add_style('ChapterNumber', WD_STYLE_TYPE.PARAGRAPH)
        except:
            chapter_num_style = styles['ChapterNumber']
        chapter_num_style.font.name = 'Arial Black'
        chapter_num_style.font.size = Pt(48)
        chapter_num_style.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        chapter_num_style.font.bold = True
        chapter_num_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_num_style.paragraph_format.space_after = Pt(12)

        # Chapter Title Style
        try:
            chapter_title_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            chapter_title_style = styles['ChapterTitle']
        chapter_title_style.font.name = 'Arial'
        chapter_title_style.font.size = Pt(24)
        chapter_title_style.font.color.rgb = RGBColor(0, 0, 0)
        chapter_title_style.font.bold = True
        chapter_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_title_style.paragraph_format.space_after = Pt(36)
        chapter_title_style.paragraph_format.keep_with_next = True

        # Main Heading Style (like Harrison's section headings)
        try:
            heading_style = styles.add_style('MainHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            heading_style = styles['MainHeading']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        heading_style.font.bold = True
        heading_style.font.all_caps = True
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
        heading_style.paragraph_format.keep_with_next = True

        # Subheading Style
        try:
            subheading_style = styles.add_style('SubHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subheading_style = styles['SubHeading']
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(12)
        subheading_style.font.color.rgb = RGBColor(0, 0, 139)
        subheading_style.font.bold = True
        subheading_style.font.italic = True
        subheading_style.paragraph_format.space_before = Pt(18)
        subheading_style.paragraph_format.space_after = Pt(6)

        # Body Text Style (Harrison's uses serif font)
        try:
            body_style = styles.add_style('HarrisonBody', WD_STYLE_TYPE.PARAGRAPH)
        except:
            body_style = styles['HarrisonBody']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(11)
        body_style.font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.space_after = Pt(8)
        body_style.paragraph_format.first_line_indent = Cm(0.5)

        # Clinical Pearl Box Style
        try:
            pearl_style = styles.add_style('ClinicalPearl', WD_STYLE_TYPE.PARAGRAPH)
        except:
            pearl_style = styles['ClinicalPearl']
        pearl_style.font.name = 'Arial'
        pearl_style.font.size = Pt(10)
        pearl_style.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
        pearl_style.paragraph_format.left_indent = Cm(0.5)
        pearl_style.paragraph_format.right_indent = Cm(0.5)
        pearl_style.paragraph_format.space_before = Pt(6)
        pearl_style.paragraph_format.space_after = Pt(6)

        # Table/Figure Caption Style
        try:
            caption_style = styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
        except:
            caption_style = styles['FigureCaption']
        caption_style.font.name = 'Arial'
        caption_style.font.size = Pt(9)
        caption_style.font.bold = True
        caption_style.paragraph_format.space_before = Pt(6)
        caption_style.paragraph_format.space_after = Pt(12)

        # Key Points Style
        try:
            keypoint_style = styles.add_style('KeyPoint', WD_STYLE_TYPE.PARAGRAPH)
        except:
            keypoint_style = styles['KeyPoint']
        keypoint_style.font.name = 'Arial'
        keypoint_style.font.size = Pt(11)
        keypoint_style.font.bold = True
        keypoint_style.font.color.rgb = RGBColor(128, 0, 0)  # Maroon
        keypoint_style.paragraph_format.left_indent = Cm(0.5)
        keypoint_style.paragraph_format.space_before = Pt(6)
        keypoint_style.paragraph_format.space_after = Pt(6)

    def add_chapter(self, title, number=None):
        """Add a new chapter with Harrison's style formatting"""
        # Start new page for each chapter
        if self.chapter_number > 0:
            self.doc.add_page_break()

        self.chapter_number += 1

        # Chapter number (e.g., "CHAPTER 1")
        chapter_num = self.doc.add_paragraph()
        chapter_num.style = 'ChapterNumber'
        chapter_num.add_run(f"CHAPTER {number if number else self.chapter_number}")

        # Chapter title
        chapter_title = self.doc.add_paragraph()
        chapter_title.style = 'ChapterTitle'
        chapter_title.add_run(title.upper())

        # Horizontal line under title
        self.add_horizontal_line()

        # Add some space
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    def add_horizontal_line(self):
        """Add a professional horizontal line"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        # Add horizontal line using borders
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_main_heading(self, text):
        """Add main section heading (Harrison's style - all caps)"""
        heading = self.doc.add_paragraph()
        heading.style = 'MainHeading'
        heading.add_run(text.upper())

    def add_subheading(self, text):
        """Add subheading"""
        subheading = self.doc.add_paragraph()
        subheading.style = 'SubHeading'
        subheading.add_run(text)

    def add_body_text(self, text):
        """Add body text with proper formatting"""
        # Split into paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = self.doc.add_paragraph()
                p.style = 'HarrisonBody'
                p.add_run(para.strip())

    def add_clinical_box(self, title, content, box_type='pearl'):
        """Add a clinical box (Pearl, Key Point, etc.)"""
        # Create table for box effect
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Set cell properties
        cell = table.cell(0, 0)
        cell_xml = cell._element
        cell_props = cell_xml.get_or_add_tcPr()

        # Add shading based on type
        shading = OxmlElement('w:shd')
        if box_type == 'pearl':
            shading.set(qn('w:fill'), 'E8F5E9')  # Light green
        elif box_type == 'warning':
            shading.set(qn('w:fill'), 'FFF3E0')  # Light orange
        else:
            shading.set(qn('w:fill'), 'E3F2FD')  # Light blue
        cell_props.append(shading)

        # Add title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"▶ {title}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = 'Arial'

        # Add content
        content_para = cell.add_paragraph()
        content_para.style = 'ClinicalPearl'
        content_para.add_run(content)

        # Add spacing after box
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    def add_key_points_list(self, points):
        """Add a list of key points"""
        self.add_clinical_box("KEY POINTS", "", "warning")

        for i, point in enumerate(points, 1):
            p = self.doc.add_paragraph()
            p.style = 'KeyPoint'
            p.add_run(f"{i}. {point}")

    def add_definition(self, term, definition):
        """Add a definition in Harrison's style"""
        p = self.doc.add_paragraph()
        p.style = 'HarrisonBody'

        # Term in bold
        term_run = p.add_run(f"{term}: ")
        term_run.font.bold = True

        # Definition in regular text
        def_run = p.add_run(definition)

    def add_table_or_figure(self, title, caption, image_data=None):
        """Add a table or figure with proper caption"""
        # Caption
        caption_para = self.doc.add_paragraph()
        caption_para.style = 'FigureCaption'
        caption_para.add_run(title)

        # Image if available
        if image_data:
            try:
                img_bytes = base64.b64decode(image_data)
                img_stream = BytesIO(img_bytes)
                self.doc.add_picture(img_stream, width=Inches(5.5))
            except:
                self.doc.add_paragraph("[Figure not available]")

        # Description
        if caption:
            desc_para = self.doc.add_paragraph()
            desc_para.style = 'FigureCaption'
            desc_para.add_run(caption)
            desc_para.runs[0].font.bold = False
            desc_para.runs[0].font.italic = True

        # Add spacing
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    def add_header_footer(self, chapter_title):
        """Add headers and footers to all sections"""
        for section in self.doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"Chapter {self.chapter_number}: {chapter_title}"
            header_para.style.font.size = Pt(10)
            header_para.style.font.italic = True
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer with page number
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Note: Actual page numbers require fields which python-docx doesn't fully support


def create_harrison_style_document(data):
    """Create a Harrison's style educational document"""
    hd = HarrisonStyleDocument()

    # Title page (like a textbook cover page)
    title_para = hd.doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)

    # Main title
    title_run = title_para.add_run("CRITICAL APPRAISAL IN MEDICINE")
    title_run.font.name = 'Arial Black'
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(139, 0, 0)
    title_run.font.bold = True

    # Subtitle
    subtitle_para = hd.doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(24)
    subtitle_run = subtitle_para.add_run("A Comprehensive Educational Guide")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.italic = True

    # Paper title
    paper_para = hd.doc.add_paragraph()
    paper_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paper_para.paragraph_format.space_before = Pt(48)
    paper_para.add_run(sanitize_text(data.get('title', ''), 'title'))
    paper_para.runs[0].font.size = Pt(14)

    # Add page break after title page
    hd.doc.add_page_break()

    # Map sections to chapters
    chapter_mapping = {
        'Paper Basics': 1,
        'Executive Summary': 2,
        'Study Design': 3,
        'Historical Context & Impact': 4,
        'Critical Appraisal Checklist': 5,
        'Methods Deep-Dive': 6,
        'How This Study Was Done': 7,
        'Results Interpretation': 8,
        'Tables and Figures': 9,
        'Critical Analysis': 10,
        'Strengths and Weaknesses': 11,
        'Clinical Implications': 12,
        'Study Impact': 13,
        'Defense Questions & Answers': 14,
        'Funding or Conflict of Interest': 15,
        'Related Research from Web': 16
    }

    # Process sections as chapters
    for section in data.get('sections', []):
        heading = sanitize_text(section.get('heading', ''))
        content = sanitize_text(section.get('content', ''))

        # Get chapter number
        chapter_num = None
        for key, num in chapter_mapping.items():
            if key in heading or heading in key:
                chapter_num = num
                break

        # Add chapter
        hd.add_chapter(heading, chapter_num)

        # Add introduction paragraph if content starts directly
        if content and not content.startswith('\n'):
            intro_para = hd.doc.add_paragraph()
            intro_para.style = 'HarrisonBody'
            # Extract first paragraph as introduction
            first_para = content.split('\n\n')[0]
            intro_para.add_run(first_para)
            intro_para.runs[0].font.italic = True
            content = '\n\n'.join(content.split('\n\n')[1:])

        # Process content sections
        if content:
            # Look for subsections
            subsections = re.split(r'\n(?=[A-Z][A-Za-z\s]+:)', content)

            for subsection in subsections:
                if ':' in subsection and subsection.index(':') < 100:
                    parts = subsection.split(':', 1)
                    if len(parts) == 2:
                        hd.add_subheading(parts[0].strip())
                        hd.add_body_text(parts[1].strip())
                else:
                    hd.add_body_text(subsection)

        # Add term definitions if present
        if section.get('term_definitions'):
            hd.add_main_heading("KEY TERMINOLOGY")
            for term, definition in section['term_definitions'].items():
                hd.add_definition(sanitize_text(term), sanitize_text(definition))

        # Add teaching notes as clinical pearls
        if section.get('teaching_notes'):
            hd.add_main_heading("CLINICAL PEARLS")
            for note in section['teaching_notes']:
                hd.add_clinical_box(
                    "Clinical Pearl",
                    sanitize_text(note),
                    'pearl'
                )

        # Add checklist items
        if section.get('checklist_items'):
            hd.add_main_heading("ASSESSMENT CHECKLIST")
            checklist_points = []
            for item in section['checklist_items']:
                assessment = item.get('assessment', 'Unclear')
                symbol = '✓' if assessment.lower() in ['present', 'yes'] else '✗'
                text = f"{symbol} {item.get('item', '')} - {item.get('rationale', '')}"
                checklist_points.append(sanitize_text(text))

            if checklist_points:
                hd.add_key_points_list(checklist_points)

        # Add tables and figures
        if section.get('images'):
            hd.add_main_heading("VISUAL ELEMENTS")
            for img in section['images']:
                hd.add_table_or_figure(
                    sanitize_text(img.get('title', 'Figure')),
                    sanitize_text(img.get('explanation', '')),
                    img.get('base64')
                )

    # Add references/metadata as final chapter
    if data.get('metadata'):
        hd.add_chapter("REFERENCES AND METADATA", 99)
        meta = data['metadata']

        if meta.get('doi'):
            hd.add_definition("DOI", meta['doi'])
        if meta.get('journal'):
            hd.add_definition("Journal", meta['journal'])
        if meta.get('year'):
            hd.add_definition("Publication Year", str(meta['year']))
        if meta.get('authors'):
            hd.add_definition("Authors", ', '.join(meta.get('authors', [])))

    return hd.doc


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python educational_word_generator_harrison.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_json = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load input data
        with open(input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate Harrison's style document
        doc = create_harrison_style_document(data)

        # Save to file
        doc.save(output_path)

        print(json.dumps({
            "success": True,
            "message": f"Harrison's style educational document created: {output_path}"
        }))

    except Exception as e:
        import traceback
        print(json.dumps({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }))
        sys.exit(1)


if __name__ == "__main__":
    main()