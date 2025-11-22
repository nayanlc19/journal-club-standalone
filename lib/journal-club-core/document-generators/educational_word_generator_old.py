"""
Educational Word Document Generator
Creates comprehensive, teaching-focused Word documents with clean formatting
Replaces the PDF generator for better editability
"""

import sys
import json
import base64
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Add the lib directory to path to import universal sanitizer
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

try:
    # Try to import the universal sanitizer
    from universal_text_sanitizer import sanitize_text, detect_publisher
except ImportError:
    # Fallback to basic sanitizer if universal is not available
    def sanitize_text(text, context='general'):
        """Fallback sanitizer if universal is not available"""
        if not text:
            return text

        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

        # Replace special Unicode spaces with regular spaces
        text = text.replace('\u202f', ' ')  # Narrow no-break space
        text = text.replace('\u200b', '')   # Zero-width space
        text = text.replace('\u00a0', ' ')  # Non-breaking space
        text = text.replace('\ufeff', '')   # Zero-width no-break space

        return text

    def detect_publisher(text):
        return 'unknown'

def parse_markdown_to_text(text):
    """
    Convert markdown to plain text for Word (Word handles formatting differently)
    """
    if not text:
        return ""

    # Sanitize first
    text = sanitize_text(text)

    # Remove emoji characters
    text = re.sub(r'[🔵🟢📋📚💡✅🎓📖🟡🔴⚠️❌✨🚀📄⏱️🎯🎉■●◆]', '', text)

    # Keep markdown as-is for now, Word will render it better than trying to convert
    return text


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style as hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def create_educational_document(data):
    """
    Create a comprehensive educational Word document

    Args:
        data: {
            "title": "Study Title",
            "metadata": {
                "authors": "...",
                "journal": "...",
                "year": "..."
            },
            "sections": [
                {
                    "heading": "Section Title",
                    "content": "Detailed content",
                    "explanations": {
                        "term1": "Detailed explanation...",
                        "term2": "Another explanation..."
                    },
                    "teaching_notes": ["Note 1", "Note 2"],
                    "checklist_items": [
                        {
                            "item": "Checklist item",
                            "rationale": "Why this matters",
                            "assessment": "Present/Absent/Unclear"
                        }
                    ],
                    "images": [...]
                }
            ]
        }
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title (sanitized)
    title = doc.add_heading(sanitize_text(data['title']), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(26, 54, 93)

    # Metadata
    if data.get('metadata'):
        meta = data['metadata']
        if meta.get('authors'):
            p = doc.add_paragraph()
            p.add_run('Authors: ').bold = True
            p.add_run(meta['authors'])

        if meta.get('journal'):
            p = doc.add_paragraph()
            p.add_run('Journal: ').bold = True
            run = p.add_run(meta['journal'])
            run.italic = True

        if meta.get('year'):
            p = doc.add_paragraph()
            p.add_run('Year: ').bold = True
            p.add_run(meta['year'])

    doc.add_paragraph()  # Spacing

    # Introduction box
    intro = doc.add_paragraph()
    intro_run = intro.add_run('How to Use This Document')
    intro_run.bold = True
    intro_run.font.size = Pt(12)
    intro_run.font.color.rgb = RGBColor(44, 82, 130)

    doc.add_paragraph(
        'This comprehensive guide is designed to help you understand EVERY aspect of this research paper - '
        'from the simplest concepts to the most complex methodology. No prior knowledge assumed. '
        'Each term is explained, each checklist justified, each decision analyzed.'
    )

    color_code = doc.add_paragraph()
    color_code.add_run('Color Code:').bold = True
    doc.add_paragraph('• Blue text = Term explanations', style='List Bullet')
    doc.add_paragraph('• Green text = Teaching notes', style='List Bullet')
    doc.add_paragraph('• Tables = Checklists with rationale', style='List Bullet')

    doc.add_page_break()

    # Process sections
    for section in data.get('sections', []):
        # Section heading
        heading = doc.add_heading(section['heading'], 1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(44, 82, 130)

        # Main content
        if section.get('content'):
            paragraphs = section['content'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Parse markdown formatting
                    para_text = parse_markdown_to_text(para.strip())

                    # Handle bold **text**
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*[^*]+\*\*)', para_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

        # Term Explanations
        if section.get('explanations'):
            heading = doc.add_heading('Term Explanations:', 2)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(45, 121, 199)

            for term, explanation in section['explanations'].items():
                p = doc.add_paragraph()
                term_run = p.add_run(f"{term}: ")
                term_run.bold = True
                term_run.font.color.rgb = RGBColor(45, 121, 199)
                p.add_run(parse_markdown_to_text(explanation))

        # Teaching Notes
        if section.get('teaching_notes'):
            heading = doc.add_heading('Teaching Notes:', 2)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(56, 161, 105)

            for note in section['teaching_notes']:
                p = doc.add_paragraph(parse_markdown_to_text(note), style='List Bullet')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(56, 161, 105)

        # Checklist Items
        if section.get('checklist_items'):
            doc.add_heading('Critical Appraisal Checklist:', 2)

            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Why This Matters'
            header_cells[2].text = 'Assessment'

            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data rows
            for item in section['checklist_items']:
                row_cells = table.add_row().cells
                row_cells[0].text = parse_markdown_to_text(item['item'])
                row_cells[1].text = parse_markdown_to_text(item['rationale'])
                row_cells[2].text = parse_markdown_to_text(item['assessment'])

            doc.add_paragraph()  # Spacing

        # Images with explanations
        for image_data in section.get('images', []):
            # Image title
            img_heading = doc.add_heading(parse_markdown_to_text(image_data['title']), 3)

            # Insert image if available
            if image_data.get('base64'):
                try:
                    img_bytes = base64.b64decode(image_data['base64'])
                    img_stream = BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Refer to original paper for this visual element]")
            else:
                # No image data
                page_num = image_data.get('pageNumber', '?')
                doc.add_paragraph(f"[See original paper - page {page_num}]")

            # Image explanation
            if image_data.get('explanation'):
                p = doc.add_paragraph()
                exp_run = p.add_run(parse_markdown_to_text(image_data['explanation']))
                exp_run.italic = True
                exp_run.font.size = Pt(10)
                exp_run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()  # Spacing

    return doc


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python educational_word_generator.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_json = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load input data
        with open(input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate document
        doc = create_educational_document(data)

        # Save to file
        doc.save(output_path)

        print(json.dumps({
            "success": True,
            "output": output_path,
            "message": "Educational Word document created successfully"
        }))

    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": str(e)
        }))
        sys.exit(1)


if __name__ == '__main__':
    main()
