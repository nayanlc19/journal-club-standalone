"""
Perfect Harrison's Textbook Style Educational Document Generator
Creates professional medical textbook-style documents with proper tables and formatting
Inspired by Harrison's Principles of Internal Medicine - Fixed Version
"""

import sys
import json
import base64
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Add the lib directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

try:
    from universal_text_sanitizer import sanitize_text, detect_publisher
except ImportError:
    def sanitize_text(text, context='general'):
        if not text:
            return text
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        text = text.replace('\u202f', ' ').replace('\u200b', '').replace('\u00a0', ' ')
        return text
    def detect_publisher(text):
        return 'unknown'


class PerfectHarrisonDocument:
    """Create a perfect Harrison's Textbook of Medicine style document"""

    def __init__(self):
        self.doc = Document()
        self.setup_page_layout()
        self.setup_styles()
        self.chapter_number = 0

    def setup_page_layout(self):
        """Setup page layout for medical textbook style"""
        sections = self.doc.sections
        for section in sections:
            # Standard textbook margins
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            # Page size (Letter)
            section.page_height = Cm(27.94)
            section.page_width = Cm(21.59)

            # Headers and footers space
            section.header_distance = Cm(1.27)
            section.footer_distance = Cm(1.27)

    def setup_styles(self):
        """Setup Harrison's-style formatting"""
        styles = self.doc.styles

        # Chapter Number Style
        try:
            chapter_num_style = styles.add_style('ChapterNumber', WD_STYLE_TYPE.PARAGRAPH)
        except:
            chapter_num_style = styles['ChapterNumber']
        chapter_num_style.font.name = 'Arial Black'
        chapter_num_style.font.size = Pt(48)
        chapter_num_style.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        chapter_num_style.font.bold = True
        chapter_num_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_num_style.paragraph_format.space_after = Pt(18)

        # Chapter Title Style
        try:
            chapter_title_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            chapter_title_style = styles['ChapterTitle']
        chapter_title_style.font.name = 'Arial'
        chapter_title_style.font.size = Pt(24)
        chapter_title_style.font.color.rgb = RGBColor(0, 0, 0)
        chapter_title_style.font.bold = True
        chapter_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_title_style.paragraph_format.space_after = Pt(24)
        chapter_title_style.paragraph_format.keep_with_next = True

        # Main Heading Style
        try:
            heading_style = styles.add_style('MainHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            heading_style = styles['MainHeading']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        heading_style.font.bold = True
        heading_style.font.all_caps = True
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(12)
        heading_style.paragraph_format.keep_with_next = True

        # Subheading Style
        try:
            subheading_style = styles.add_style('SubHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subheading_style = styles['SubHeading']
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(12)
        subheading_style.font.color.rgb = RGBColor(0, 0, 139)
        subheading_style.font.bold = True
        subheading_style.font.italic = True
        subheading_style.paragraph_format.space_before = Pt(12)
        subheading_style.paragraph_format.space_after = Pt(6)

        # Body Text Style
        try:
            body_style = styles.add_style('HarrisonBody', WD_STYLE_TYPE.PARAGRAPH)
        except:
            body_style = styles['HarrisonBody']
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(11)
        body_style.font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.first_line_indent = Cm(0.5)

    def add_chapter(self, title, number=None):
        """Add a new chapter with Harrison's style formatting"""
        # Start new page for each chapter
        if self.chapter_number > 0:
            self.doc.add_page_break()

        self.chapter_number += 1

        # Chapter number
        chapter_num = self.doc.add_paragraph()
        chapter_num.style = 'ChapterNumber'
        chapter_num.add_run(f"CHAPTER {number if number else self.chapter_number}")

        # Chapter title
        chapter_title = self.doc.add_paragraph()
        chapter_title.style = 'ChapterTitle'
        chapter_title.add_run(title.upper())

        # Horizontal line under title
        self.add_horizontal_line()

    def add_horizontal_line(self):
        """Add a professional horizontal line"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)

        # Add horizontal line using borders
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_main_heading(self, text):
        """Add main section heading"""
        heading = self.doc.add_paragraph()
        heading.style = 'MainHeading'
        heading.add_run(text.upper())

    def add_subheading(self, text):
        """Add subheading"""
        subheading = self.doc.add_paragraph()
        subheading.style = 'SubHeading'
        subheading.add_run(text)

    def add_body_text(self, text):
        """Add body text with proper formatting"""
        if not text or not text.strip():
            return

        # Clean up the text
        text = sanitize_text(text)

        # Split into paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if para:
                p = self.doc.add_paragraph()
                p.style = 'HarrisonBody'
                p.add_run(para)

    def parse_table_data(self, text):
        """Parse text to determine if it should be a table"""
        # Look for patterns that indicate table data
        lines = text.strip().split('\n')

        # Check for timeline/year patterns
        if any(re.match(r'^\d{4}', line.strip()) for line in lines):
            return self.parse_timeline_table(lines)

        # Check for statistical data patterns
        if any(re.search(r'\d+\s*/\s*\d+|[\d.]+\s*%|\d+\.\d+', line) for line in lines):
            return self.parse_stats_table(lines)

        # Check for structured list with colons
        if len(lines) > 2 and all(':' in line for line in lines[:3]):
            return self.parse_definition_table(lines)

        return None

    def parse_timeline_table(self, lines):
        """Parse timeline data into table format"""
        table_data = []
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Look for year pattern
            match = re.match(r'^(\d{4})\s+(.+)', line)
            if match:
                year = match.group(1)
                description = match.group(2)
                table_data.append([year, description])
            elif table_data:  # Continuation of previous row
                table_data[-1][1] += ' ' + line

        return table_data if table_data else None

    def parse_stats_table(self, lines):
        """Parse statistical data into table format"""
        table_data = []
        headers = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Try to split by common delimiters
            parts = re.split(r'\s{2,}|\t', line)
            if len(parts) > 1:
                # Check if it looks like headers (no numbers)
                if not any(re.search(r'\d', part) for part in parts):
                    headers = parts
                else:
                    table_data.append(parts)

        if headers and table_data:
            return [headers] + table_data
        elif table_data:
            return table_data

        return None

    def parse_definition_table(self, lines):
        """Parse definition/colon-separated data into table format"""
        table_data = []
        for line in lines:
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    table_data.append([parts[0].strip(), parts[1].strip()])

        return table_data if table_data else None

    def create_table(self, data, headers=None, style='medium'):
        """Create a properly formatted Word table"""
        if not data:
            return

        # Determine number of columns
        num_cols = max(len(row) for row in data)
        num_rows = len(data) + (1 if headers else 0)

        # Create table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Apply table style
        if style == 'medium':
            table.style = 'Medium Shading 1 Accent 1'
        elif style == 'light':
            table.style = 'Light Grid Accent 1'
        else:
            table.style = 'Table Grid'

        # Add headers if provided
        row_offset = 0
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers[:num_cols]):
                header_cells[i].text = str(header)
                # Bold headers
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            row_offset = 1

        # Add data
        for i, row_data in enumerate(data):
            cells = table.rows[i + row_offset].cells
            for j, cell_data in enumerate(row_data[:num_cols]):
                cells[j].text = str(cell_data) if cell_data else ''
                # Set font for cells
                for paragraph in cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'

        # Adjust column widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(6.5 / num_cols)

        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    def add_clinical_box(self, title, content, box_type='pearl'):
        """Add a clinical box with proper formatting"""
        # Create table for box effect
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set width
        table.width = Inches(6)

        # Set cell properties
        cell = table.cell(0, 0)
        cell_xml = cell._element
        cell_props = cell_xml.get_or_add_tcPr()

        # Add shading based on type
        shading = OxmlElement('w:shd')
        if box_type == 'pearl':
            shading.set(qn('w:fill'), 'E8F5E9')  # Light green
        elif box_type == 'warning':
            shading.set(qn('w:fill'), 'FFF3E0')  # Light orange
        else:
            shading.set(qn('w:fill'), 'E3F2FD')  # Light blue
        cell_props.append(shading)

        # Add title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"▶ {title}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = 'Arial'
        title_run.font.color.rgb = RGBColor(0, 0, 139)

        # Add content
        if content:
            content_para = cell.add_paragraph()
            content_para.add_run(content)
            content_para.runs[0].font.size = Pt(10)
            content_para.runs[0].font.name = 'Times New Roman'

        # Add spacing after box
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    def add_statistical_results(self, content):
        """Parse and format statistical results into tables"""
        import re
        lines = content.strip().split('\n')

        # Collect statistical data for table
        statistical_data = []
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line contains statistical data
            has_stats = any(indicator in line for indicator in
                          ['%', 'CI:', 'HR:', 'RR:', 'OR:', 'p=', 'P=', 'n=', 'vs', 'versus'])

            if has_stats:
                # Parse the line to extract outcome and values
                if ':' in line:
                    parts = line.split(':', 1)
                    outcome = parts[0].strip()
                    value = parts[1].strip() if len(parts) > 1 else ''

                    # Clean up outcome name
                    outcome = outcome.replace('•', '').replace('-', '').strip()

                    # Try to extract treatment vs control values
                    if 'vs' in value or 'versus' in value:
                        value_parts = re.split(r'\s+(?:vs|versus)\s+', value)
                        if len(value_parts) == 2:
                            statistical_data.append([outcome, value_parts[0].strip(), value_parts[1].strip()])
                        else:
                            statistical_data.append([outcome, value, ''])
                    else:
                        statistical_data.append([outcome, value, ''])
                elif '%' in line or 'CI' in line:
                    # Try to extract percentage or CI data
                    statistical_data.append(['', line, ''])

        # Create table if we have statistical data
        if statistical_data:
            # Add table title
            table_title = self.doc.add_paragraph()
            table_title.add_run("Statistical Results")
            table_title.runs[0].font.bold = True
            table_title.runs[0].font.size = Pt(12)
            table_title.paragraph_format.space_before = Pt(12)
            table_title.paragraph_format.space_after = Pt(6)

            # Create the table
            headers = ['Outcome', 'Treatment/Value', 'Control/Comparison']
            self.create_table(statistical_data, headers=headers, style='colorful')
        else:
            # Fall back to paragraph format if no clear statistical data
            for line in lines:
                if line.strip():
                    p = self.doc.add_paragraph()
                    p.style = 'HarrisonBody'
                    p.add_run(line)


def create_perfect_harrison_document(data):
    """Create a perfect Harrison's style educational document"""
    hd = PerfectHarrisonDocument()

    # Title page
    title_para = hd.doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)

    # Main title
    title_run = title_para.add_run("CRITICAL APPRAISAL IN MEDICINE")
    title_run.font.name = 'Arial Black'
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(139, 0, 0)
    title_run.font.bold = True

    # Subtitle
    subtitle_para = hd.doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(24)
    subtitle_run = subtitle_para.add_run("A Comprehensive Educational Guide")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.italic = True

    # Paper title
    paper_para = hd.doc.add_paragraph()
    paper_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paper_para.paragraph_format.space_before = Pt(48)
    paper_para.add_run(sanitize_text(data.get('title', ''), 'title'))
    paper_para.runs[0].font.size = Pt(14)

    # Add page break after title page
    hd.doc.add_page_break()

    # Table of Contents
    hd.add_main_heading("TABLE OF CONTENTS")

    toc_items = [
        "1. Paper Basics",
        "2. Executive Summary",
        "3. Study Design",
        "4. Historical Context & Impact",
        "5. Critical Appraisal Checklist",
        "6. Methods Deep-Dive",
        "7. How This Study Was Done",
        "8. Results Interpretation",
        "9. Tables and Figures",
        "10. Critical Analysis",
        "11. Strengths and Weaknesses",
        "12. Clinical Implications",
        "13. Study Impact",
        "14. Defense Questions & Answers",
        "15. Funding and Conflicts",
        "16. Related Research"
    ]

    for item in toc_items:
        p = hd.doc.add_paragraph()
        p.style = 'HarrisonBody'
        p.add_run(item)
        p.runs[0].font.bold = True

    hd.doc.add_page_break()

    # Process sections as chapters
    chapter_mapping = {
        'Paper Basics': 1,
        'Executive Summary': 2,
        'Study Design': 3,
        'Historical Context': 4,
        'Critical Appraisal Checklist': 5,
        'Methods Deep-Dive': 6,
        'How This Study Was Done': 7,
        'Results Interpretation': 8,
        'Tables and Figures': 9,
        'Critical Analysis': 10,
        'Strengths': 11,
        'Weaknesses': 11,
        'Clinical Implications': 12,
        'Study Impact': 13,
        'Defense Questions': 14,
        'Funding': 15,
        'Related Research': 16
    }

    for section in data.get('sections', []):
        heading = sanitize_text(section.get('heading', ''))
        content = sanitize_text(section.get('content', ''))

        # Get chapter number
        chapter_num = None
        for key, num in chapter_mapping.items():
            if key.lower() in heading.lower():
                chapter_num = num
                break

        # Add chapter
        hd.add_chapter(heading, chapter_num)

        # Process content intelligently
        if content:
            # Check if content should be a table
            table_data = hd.parse_table_data(content)

            if table_data:
                # Create appropriate table
                if 'timeline' in heading.lower() or 'evolution' in heading.lower() or 'history' in heading.lower():
                    hd.add_main_heading("TIMELINE")
                    hd.create_table(table_data, headers=['Year', 'Development/Milestone', 'Significance/Impact'], style='medium')
                elif 'statistic' in heading.lower() or 'result' in heading.lower():
                    hd.add_main_heading("STATISTICAL RESULTS")
                    hd.create_table(table_data, style='light')
                else:
                    hd.create_table(table_data)
            else:
                # Process as regular text with proper formatting
                # Split content into logical sections
                sections = re.split(r'\n(?=[A-Z][A-Za-z\s]+:|\d+\.)', content)

                for section_text in sections:
                    section_text = section_text.strip()
                    if not section_text:
                        continue

                    # Check for numbered lists
                    if re.match(r'^\d+\.', section_text):
                        # Format as numbered paragraph
                        hd.add_body_text(section_text)
                    # Check for subsection headers (text followed by colon)
                    elif ':' in section_text and section_text.index(':') < 100:
                        parts = section_text.split(':', 1)
                        if len(parts) == 2 and len(parts[0]) < 100:
                            hd.add_subheading(parts[0].strip())

                            # Check if the content after colon should be a table
                            remaining = parts[1].strip()
                            table_data = hd.parse_table_data(remaining)
                            if table_data:
                                hd.create_table(table_data)
                            else:
                                hd.add_body_text(remaining)
                        else:
                            hd.add_body_text(section_text)
                    else:
                        # Check if this text contains statistics that should be a table
                        if any(indicator in section_text for indicator in ['%', 'CI:', 'HR:', 'vs', 'versus', 'P=', 'p=']):
                            hd.add_statistical_results(section_text)
                        else:
                            hd.add_body_text(section_text)

        # Process statistical results specially
        if 'result' in heading.lower() or 'outcome' in heading.lower() or 'finding' in heading.lower():
            # Always try to create tables for results sections
            hd.add_statistical_results(content)

        # Add term definitions as a proper table
        if section.get('term_definitions'):
            hd.add_main_heading("KEY TERMINOLOGY")
            term_data = [[term, definition] for term, definition in section['term_definitions'].items()]
            hd.create_table(term_data, headers=['Term', 'Definition'], style='light')

        # Add teaching notes as clinical pearls
        if section.get('teaching_notes'):
            hd.add_main_heading("CLINICAL PEARLS")
            for note in section['teaching_notes']:
                hd.add_clinical_box("Clinical Pearl", sanitize_text(note), 'pearl')

        # Add checklist items as a table
        if section.get('checklist_items'):
            hd.add_main_heading("ASSESSMENT CHECKLIST")
            checklist_data = []
            for item in section['checklist_items']:
                assessment = item.get('assessment', 'Unclear')
                symbol = '✓' if assessment.lower() in ['present', 'yes', 'adequate'] else '✗'
                checklist_data.append([
                    symbol,
                    item.get('item', ''),
                    item.get('rationale', ''),
                    assessment
                ])

            hd.create_table(
                checklist_data,
                headers=['Status', 'Criterion', 'Rationale', 'Assessment'],
                style='medium'
            )

        # Add images properly
        if section.get('images'):
            hd.add_main_heading("VISUAL ELEMENTS")
            for i, img in enumerate(section['images'], 1):
                # Add figure title
                fig_title = hd.doc.add_paragraph()
                fig_title.add_run(f"Figure {i}. {sanitize_text(img.get('title', 'Untitled'))}")
                fig_title.runs[0].font.bold = True
                fig_title.runs[0].font.size = Pt(11)
                fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add image if available
                if img.get('base64'):
                    try:
                        img_bytes = base64.b64decode(img['base64'])
                        img_stream = BytesIO(img_bytes)
                        picture = hd.doc.add_picture(img_stream, width=Inches(5.5))
                        # Center the image
                        last_paragraph = hd.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        hd.doc.add_paragraph("[Figure could not be displayed]")

                # Add caption
                if img.get('explanation'):
                    caption = hd.doc.add_paragraph()
                    caption.add_run(sanitize_text(img['explanation']))
                    caption.runs[0].font.size = Pt(10)
                    caption.runs[0].font.italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(18)

    # Add references/metadata as final chapter
    if data.get('metadata'):
        hd.add_chapter("REFERENCES AND METADATA", 99)
        meta = data['metadata']

        metadata_table = []
        if meta.get('doi'):
            metadata_table.append(['DOI', meta['doi']])
        if meta.get('journal'):
            metadata_table.append(['Journal', meta['journal']])
        if meta.get('year'):
            metadata_table.append(['Publication Year', str(meta['year'])])
        if meta.get('authors'):
            metadata_table.append(['Authors', ', '.join(meta.get('authors', []))])

        if metadata_table:
            hd.create_table(metadata_table, headers=['Field', 'Value'], style='light')

    return hd.doc


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python educational_word_generator_perfect.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_json = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load input data
        with open(input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate perfect Harrison's style document
        doc = create_perfect_harrison_document(data)

        # Save to file
        doc.save(output_path)

        print(json.dumps({
            "success": True,
            "message": f"Perfect Harrison's style document created: {output_path}"
        }))

    except Exception as e:
        import traceback
        print(json.dumps({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }))
        sys.exit(1)


if __name__ == "__main__":
    main()