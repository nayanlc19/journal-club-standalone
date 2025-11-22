#!/usr/bin/env python3
"""
Educational Word Generator V2 - With Professional Book-Style Formatting
Replaces the existing educational_word_generator.py with enhanced formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import os
import base64
from io import BytesIO
import re

class EducationalWordGeneratorV2:
    """Generate professionally formatted educational Word documents"""

    def __init__(self):
        self.doc = Document()
        self.setup_book_styles()
        self.setup_page_layout()

    def setup_page_layout(self):
        """Configure professional page layout"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.0)

    def setup_book_styles(self):
        """Setup professional book-style formatting"""
        styles = self.doc.styles

        # Title style
        try:
            title_style = styles.add_style('MainTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Georgia'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        except:
            pass  # Style already exists

        # Chapter style
        try:
            chapter_style = styles.add_style('ChapterStyle', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Arial Black'
            chapter_style.font.size = Pt(18)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = RGBColor(128, 0, 0)
            chapter_style.paragraph_format.space_before = Pt(18)
            chapter_style.paragraph_format.space_after = Pt(12)
        except:
            pass

        # Section header
        try:
            section_style = styles.add_style('SectionHead', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Arial'
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(8)
        except:
            pass

        # Body text with proper spacing
        try:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.space_after = Pt(8)
            body_style.paragraph_format.line_spacing = 1.5
            body_style.paragraph_format.first_line_indent = Inches(0.3)
        except:
            pass

        # Result box style
        try:
            result_style = styles.add_style('ResultStyle', WD_STYLE_TYPE.PARAGRAPH)
            result_style.font.name = 'Courier New'
            result_style.font.size = Pt(10)
            result_style.font.bold = True
            result_style.paragraph_format.left_indent = Inches(0.4)
            result_style.paragraph_format.right_indent = Inches(0.4)
            result_style.paragraph_format.space_before = Pt(6)
            result_style.paragraph_format.space_after = Pt(6)
        except:
            pass

        # Interpretation style
        try:
            interp_style = styles.add_style('InterpStyle', WD_STYLE_TYPE.PARAGRAPH)
            interp_style.font.name = 'Calibri'
            interp_style.font.size = Pt(11)
            interp_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            interp_style.paragraph_format.first_line_indent = Inches(0.4)
            interp_style.paragraph_format.space_after = Pt(10)
        except:
            pass

    def add_professional_table(self, headers, rows, caption=None):
        """Create a professionally formatted table with proper borders and spacing"""

        # Add caption if provided
        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'Table: {caption}')
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.italic = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

        # Create table
        num_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'  # This adds all borders
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        table.allow_autofit = False
        total_width = Inches(6.0)
        col_width = total_width / num_cols

        # Format header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.width = col_width
            cell.text = str(header)

            # Format header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

            # Add gray background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._element.get_or_add_tcPr().append(shading)

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add data rows
        for row_data in rows:
            row = table.add_row()
            for i, value in enumerate(row_data[:num_cols]):
                cell = row.cells[i]
                cell.width = col_width

                # Handle different value types
                if isinstance(value, (int, float)):
                    cell.text = f"{value:,.1f}" if isinstance(value, float) else str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    cell.text = str(value) if value else ''
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Format cell
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        return table

    def add_result_interpretation(self, result, interpretation):
        """Add formatted result and interpretation sections"""

        # Result box with background
        p = self.doc.add_paragraph(style='ResultStyle')

        # Add shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        p._element.get_or_add_pPr().append(shading)

        run = p.add_run('Result: ')
        run.bold = True
        run.underline = True
        p.add_run(result)

        # Interpretation
        p = self.doc.add_paragraph(style='InterpStyle')
        run = p.add_run('Interpretation: ')
        run.bold = True
        run.italic = True
        p.add_run(interpretation)

    def add_section_with_formatting(self, heading, content, is_chapter=False):
        """Add a section with proper formatting"""

        if is_chapter:
            p = self.doc.add_paragraph(heading, style='ChapterStyle')
        else:
            p = self.doc.add_paragraph(heading, style='SectionHead')

        # Process content - check if it contains result/interpretation
        if isinstance(content, str):
            # Look for Result: and Interpretation: patterns
            result_pattern = r'Result:\s*(.+?)(?=Interpretation:|$)'
            interp_pattern = r'Interpretation:\s*(.+?)$'

            result_match = re.search(result_pattern, content, re.DOTALL)
            interp_match = re.search(interp_pattern, content, re.DOTALL)

            if result_match and interp_match:
                self.add_result_interpretation(
                    result_match.group(1).strip(),
                    interp_match.group(1).strip()
                )
            else:
                # Regular body text
                p = self.doc.add_paragraph(content, style='BodyText')

    def generate(self, data, output_path):
        """Generate the educational document with professional formatting"""

        # Title
        title = data.get('title', 'Critical Appraisal')
        p = self.doc.add_paragraph(title, style='MainTitle')

        # Metadata
        if 'metadata' in data:
            meta = data['metadata']
            if meta.get('authors'):
                p = self.doc.add_paragraph(meta['authors'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.italic = True
            if meta.get('journal'):
                p = self.doc.add_paragraph(f"{meta['journal']} ({meta.get('year', 'N/A')})")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(10)

        # Process sections
        chapter_num = 1
        for section in data.get('sections', []):
            heading = section.get('heading', '')
            content = section.get('content', '')

            # Check if this should be a chapter
            is_chapter = 'CHAPTER' in heading.upper() or 'RESULTS' in heading.upper()

            if is_chapter:
                self.add_section_with_formatting(f'CHAPTER {chapter_num}', '', True)
                self.add_section_with_formatting(heading.replace('CHAPTER', '').strip(), content)
                chapter_num += 1
            else:
                self.add_section_with_formatting(heading, content)

            # Add tables if present
            if 'checklist_items' in section:
                items = section['checklist_items']
                if items:
                    headers = ['Item', 'Assessment', 'Rationale']
                    rows = [[item.get('item', ''), item.get('assessment', ''), item.get('rationale', '')]
                            for item in items]
                    self.add_professional_table(headers, rows, caption=heading)

            # Add explanations if present
            if 'explanations' in section:
                self.doc.add_paragraph('Key Terms Explained', style='SectionHead')
                for term, explanation in section['explanations'].items():
                    p = self.doc.add_paragraph()
                    run = p.add_run(f'{term}: ')
                    run.bold = True
                    run.font.size = Pt(10)
                    p.add_run(explanation)
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_after = Pt(4)

            # Add teaching notes
            if 'teaching_notes' in section:
                self.doc.add_paragraph('Teaching Points', style='SectionHead')
                for note in section['teaching_notes']:
                    p = self.doc.add_paragraph(f'• {note}')
                    p.paragraph_format.left_indent = Inches(0.4)
                    p.paragraph_format.space_after = Pt(4)

        # Save document
        self.doc.save(output_path)
        return {'success': True, 'path': output_path}

def main():
    """Main function for command-line use"""

    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': 'Usage: script.py input.json output.docx'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load data
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate document
        generator = EducationalWordGeneratorV2()
        result = generator.generate(data, output_path)

        print(json.dumps(result))

    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()