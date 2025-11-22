#!/usr/bin/env python3
"""
Enhanced Book-Style Document Generator for Journal Club V2
Integrates with existing pipeline to create professionally formatted documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
import json
import os
import re
import sys
from pathlib import Path

# Import the base formatter
sys.path.append(str(Path(__file__).parent.parent.parent))

class EnhancedBookGenerator:
    """Enhanced document generator with professional book-style formatting"""

    def __init__(self):
        self.doc = Document()
        self.setup_professional_styles()
        self.setup_book_layout()

    def setup_book_layout(self):
        """Configure page for book-style layout with proper margins"""
        sections = self.doc.sections
        for section in sections:
            # Standard letter size
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)

            # Book-style margins
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.0)

            # Gutter for binding
            section.gutter = Inches(0.25)

    def setup_professional_styles(self):
        """Define comprehensive professional styles"""
        styles = self.doc.styles

        # Main title style
        title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Georgia'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(36)

        # Chapter heading style (like "CHAPTER 8")
        chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Arial Black'
        chapter_style.font.size = Pt(20)
        chapter_style.font.bold = True
        chapter_style.font.color.rgb = RGBColor(128, 0, 0)  # Dark red
        chapter_style.paragraph_format.space_before = Pt(24)
        chapter_style.paragraph_format.space_after = Pt(18)

        # Section title style
        section_style = styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Arial'
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(18)
        section_style.paragraph_format.space_after = Pt(12)

        # Subsection style
        subsection_style = styles.add_style('SubsectionTitle', WD_STYLE_TYPE.PARAGRAPH)
        subsection_style.font.name = 'Arial'
        subsection_style.font.size = Pt(12)
        subsection_style.font.bold = True
        subsection_style.font.italic = True
        subsection_style.paragraph_format.space_before = Pt(12)
        subsection_style.paragraph_format.space_after = Pt(8)

        # Professional body text
        body_style = styles.add_style('ProfessionalBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.space_after = Pt(8)
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.first_line_indent = Inches(0.3)

        # Result box style (gray background)
        result_style = styles.add_style('ResultBox', WD_STYLE_TYPE.PARAGRAPH)
        result_style.font.name = 'Courier New'
        result_style.font.size = Pt(10)
        result_style.font.bold = True
        result_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        result_style.paragraph_format.left_indent = Inches(0.5)
        result_style.paragraph_format.right_indent = Inches(0.5)
        result_style.paragraph_format.space_before = Pt(8)
        result_style.paragraph_format.space_after = Pt(8)

        # Interpretation style
        interp_style = styles.add_style('InterpretationText', WD_STYLE_TYPE.PARAGRAPH)
        interp_style.font.name = 'Calibri'
        interp_style.font.size = Pt(11)
        interp_style.font.italic = True
        interp_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        interp_style.paragraph_format.first_line_indent = Inches(0.5)
        interp_style.paragraph_format.space_before = Pt(6)
        interp_style.paragraph_format.space_after = Pt(12)

        # Bullet point style
        bullet_style = styles.add_style('BulletPoint', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Arial'
        bullet_style.font.size = Pt(10)
        bullet_style.paragraph_format.left_indent = Inches(0.5)
        bullet_style.paragraph_format.space_after = Pt(4)

    def add_formatted_table(self, data, title=None):
        """Create a professionally formatted table exactly like in medical textbooks"""

        # Add table title if provided
        if title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Table. {title}")
            run.font.bold = True
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # Determine table structure
        if isinstance(data, dict):
            # Convert dict to headers and rows
            if 'headers' in data and 'rows' in data:
                headers = data['headers']
                rows = data['rows']
            else:
                headers = list(data.keys())
                rows = [list(data.values())]
        elif isinstance(data, list) and len(data) > 0:
            # Assume first row is headers
            headers = data[0] if isinstance(data[0], list) else ['Column 1', 'Column 2']
            rows = data[1:] if len(data) > 1 else []
        else:
            return  # No valid data

        # Create table with borders
        num_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=num_cols)

        # Apply table style with full borders
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table width
        table.allow_autofit = False
        total_width = Inches(6.0)  # Standard width for readability
        col_width = total_width / num_cols

        # Format header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.width = col_width

            # Clear existing paragraphs and add formatted text
            cell.text = ''
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(header))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Arial'

            # Add gray background to header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')  # Light gray
            cell._element.get_or_add_tcPr().append(shading)

            # Center vertically
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add data rows
        for row_data in rows:
            row = table.add_row()
            for i, value in enumerate(row_data[:num_cols]):
                cell = row.cells[i]
                cell.width = col_width

                # Format cell content
                cell.text = ''
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Handle different value types
                if isinstance(value, (int, float)):
                    # Right-align numbers
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    text = f"{value:,.1f}" if isinstance(value, float) else str(value)
                else:
                    text = str(value) if value is not None else ''

                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.name = 'Arial'

                # Center vertically
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        return table

    def add_result_interpretation_block(self, result, interpretation):
        """Add a result/interpretation block like in medical textbooks"""

        # Result section with gray background
        p = self.doc.add_paragraph(style='ResultBox')

        # Add shading for result
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')  # Light gray
        p._element.get_or_add_pPr().append(shading)

        # Add "Result:" label
        run = p.add_run('Result')
        run.font.bold = True
        run.font.underline = True
        p.add_run('\n')
        p.add_run(result)

        # Interpretation section
        p = self.doc.add_paragraph(style='InterpretationText')
        run = p.add_run('Interpretation')
        run.font.bold = True
        run.font.underline = True
        run.font.italic = False
        p.add_run('\n\n')

        # Add interpretation text
        interp_run = p.add_run(interpretation)
        interp_run.font.italic = False  # Normal text for readability

    def generate_from_appraisal(self, appraisal_data, output_path, doc_type='educational'):
        """Generate a professionally formatted document from appraisal data"""

        # Extract key information
        title = appraisal_data.get('title', 'Critical Appraisal')
        pico = appraisal_data.get('clinical_question', '')

        # Add main title
        p = self.doc.add_paragraph(title, style='BookTitle')

        # Add metadata if available
        if 'authors' in appraisal_data:
            p = self.doc.add_paragraph(appraisal_data['authors'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.italic = True

        # Add CHAPTER header
        self.doc.add_paragraph('CHAPTER 8', style='ChapterHeading')
        self.doc.add_paragraph('RESULTS INTERPRETATION (SIMPLE LANGUAGE)', style='SectionTitle')

        # Clinical Question section
        if pico:
            self.doc.add_paragraph('Clinical Question (PICO)', style='SubsectionTitle')
            self.doc.add_paragraph(pico, style='ProfessionalBody')

        # Primary Outcome
        if 'primary_outcome' in appraisal_data:
            self.doc.add_paragraph('Primary Outcome', style='SubsectionTitle')
            self.doc.add_paragraph(appraisal_data['primary_outcome'], style='ProfessionalBody')

        # Results section with proper formatting
        if 'results' in appraisal_data:
            self.doc.add_paragraph('Key Results', style='SectionTitle')

            # Add result/interpretation blocks
            for result in appraisal_data.get('results', []):
                if isinstance(result, dict):
                    self.add_result_interpretation_block(
                        result.get('result', ''),
                        result.get('interpretation', '')
                    )

        # Tables section
        if 'tables' in appraisal_data:
            self.doc.add_paragraph('Statistical Analysis', style='SectionTitle')

            for table in appraisal_data['tables']:
                # Extract table data and create formatted table
                if isinstance(table, dict):
                    self.add_formatted_table(
                        table,
                        title=table.get('caption', table.get('title', ''))
                    )

        # Key Findings box
        if 'key_findings' in appraisal_data:
            self.doc.add_paragraph('Key Findings', style='SectionTitle')

            # Create a bordered box for key findings
            table = self.doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)

            # Add findings as bullet points
            for finding in appraisal_data['key_findings']:
                p = cell.add_paragraph(f"• {finding}", style='BulletPoint')

            # Light yellow background for emphasis
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'FFFACD')  # Light yellow
            cell._element.get_or_add_tcPr().append(shading)

        # Statistical Test Results
        if 'statistical_analysis' in appraisal_data:
            self.doc.add_paragraph('Statistical Tests', style='SubsectionTitle')

            stats = appraisal_data['statistical_analysis']
            if isinstance(stats, str):
                # Format p-values and confidence intervals properly
                formatted_stats = stats.replace('p<', 'p < ').replace('p=', 'p = ')
                self.doc.add_paragraph(formatted_stats, style='ProfessionalBody')
            elif isinstance(stats, dict):
                # Create a small table for statistics
                self.add_formatted_table({
                    'headers': ['Test', 'Value', 'Interpretation'],
                    'rows': [[k, v, ''] for k, v in stats.items()]
                })

        # Conclusions
        if 'conclusions' in appraisal_data:
            self.doc.add_paragraph('Conclusions', style='SectionTitle')
            self.doc.add_paragraph(appraisal_data['conclusions'], style='ProfessionalBody')

        # Critical Appraisal Score
        if 'quality_scores' in appraisal_data:
            self.doc.add_paragraph('Quality Assessment', style='SectionTitle')

            scores = appraisal_data['quality_scores']
            if isinstance(scores, dict):
                score_data = {
                    'headers': ['Criterion', 'Score', 'Percentage'],
                    'rows': []
                }
                for criterion, score in scores.items():
                    if isinstance(score, dict):
                        score_data['rows'].append([
                            criterion,
                            f"{score.get('score', 0)}/{score.get('total', 0)}",
                            f"{score.get('percentage', 0):.1f}%"
                        ])

                self.add_formatted_table(score_data, title='Study Quality Scores')

        # Save the document
        self.doc.save(output_path)
        print(f"Enhanced document saved: {output_path}")

        return output_path

def main():
    """Main function for testing or standalone use"""

    if len(sys.argv) > 1:
        # Load appraisal data from file
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            appraisal_data = json.load(f)

        output_path = sys.argv[2] if len(sys.argv) > 2 else 'enhanced_appraisal.docx'

        generator = EnhancedBookGenerator()
        generator.generate_from_appraisal(appraisal_data, output_path)
    else:
        print("Usage: python enhanced_book_generator.py <appraisal_json> [output.docx]")

        # Run test with sample data
        test_data = {
            'title': 'Critical Appraisal: Caesarean Section Study',
            'clinical_question': 'Does the Robson classification reduce CS rates?',
            'primary_outcome': 'Overall caesarean section rate',
            'results': [
                {
                    'result': '165 CS out of 352 deliveries (46.9%)',
                    'interpretation': 'Nearly half of all deliveries were by caesarean section'
                }
            ],
            'tables': [
                {
                    'caption': 'Distribution by Robson Groups',
                    'headers': ['Group', 'N', '% of CS', '% of Deliveries'],
                    'rows': [
                        ['1', 10, '6.1%', '2.8%'],
                        ['2', 33, '20.0%', '9.4%'],
                        ['5', 41, '24.8%', '11.6%']
                    ]
                }
            ],
            'key_findings': [
                'High overall CS rate of 46.9%',
                'Group 5 contributed most to CS',
                'Significant variation between groups (p < 0.001)'
            ]
        }

        generator = EnhancedBookGenerator()
        generator.generate_from_appraisal(test_data, 'test_enhanced.docx')
        print("Test document created: test_enhanced.docx")

if __name__ == "__main__":
    main()