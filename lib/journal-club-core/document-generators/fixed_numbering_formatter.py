#!/usr/bin/env python3
"""
Fixed formatter that handles numbered lists properly
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import re

def fix_numbered_content(content):
    """Fix numbered lists in content to use proper Word numbering"""

    lines = content.split('\n')
    fixed_lines = []
    in_numbered_list = False
    list_level = 0

    for line in lines:
        # Check if line starts with a number pattern (e.g., "1. ", "18. ", etc.)
        number_match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        letter_match = re.match(r'^([a-zA-Z])\.\s+(.+)', line.strip())
        bullet_match = re.match(r'^[•·▪︎▸-]\s*(.+)', line.strip())

        if number_match:
            # This is a numbered item - mark it for proper list formatting
            fixed_lines.append(f"[NUMBERED_LIST]{number_match.group(2)}")
            in_numbered_list = True
        elif letter_match:
            # This is a lettered sub-item
            fixed_lines.append(f"[LETTER_LIST]{letter_match.group(2)}")
        elif bullet_match:
            # This is a bullet point
            fixed_lines.append(f"[BULLET_LIST]{bullet_match.group(1)}")
        else:
            # Regular line - check if we need to end list
            if line.strip() == '':
                in_numbered_list = False
            fixed_lines.append(line)

    return '\n'.join(fixed_lines)

def create_document_with_proper_numbering(data, output_path):
    """Create Word document with proper numbered lists"""

    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = data.get('title', 'Critical Appraisal')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)

    # Track list state
    current_list_type = None

    # Process sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')

        # Add heading
        if heading:
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            current_list_type = None  # Reset list when new section starts

        # Process content with fixed numbering
        if content:
            # Fix numbered content
            fixed_content = fix_numbered_content(content)

            # Detect and convert tables
            cleaned_content, detected_tables = detect_and_convert_tables(fixed_content)

            # Process paragraphs
            paragraphs = cleaned_content.split('\n')

            for para in paragraphs:
                para = para.strip()
                if not para:
                    current_list_type = None  # Reset list on empty line
                    continue

                # Check for list markers
                if para.startswith('[NUMBERED_LIST]'):
                    # Use Word's built-in numbered list
                    text = para.replace('[NUMBERED_LIST]', '')
                    p = doc.add_paragraph(text, style='List Number')
                    current_list_type = 'numbered'

                elif para.startswith('[LETTER_LIST]'):
                    # Use lettered sub-list
                    text = para.replace('[LETTER_LIST]', '')
                    p = doc.add_paragraph(text, style='List Number 2')

                elif para.startswith('[BULLET_LIST]'):
                    # Use bullet list
                    text = para.replace('[BULLET_LIST]', '')
                    p = doc.add_paragraph(text, style='List Bullet')
                    current_list_type = 'bullet'

                elif para.startswith('[TABLE_') and para.endswith(']'):
                    # Table marker
                    try:
                        table_num = int(para[7:-1])
                        if table_num < len(detected_tables):
                            create_proper_table(doc, detected_tables[table_num])
                    except:
                        pass
                    current_list_type = None

                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    p.add_run(para)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(6)

                    # Only indent if not following a list
                    if current_list_type is None and not para.startswith('•'):
                        p.paragraph_format.first_line_indent = Inches(0.25)

                    current_list_type = None

        # Handle checklist items as tables
        if 'checklist_items' in section and section['checklist_items']:
            table_data = {
                'headers': ['Item', 'Assessment', 'Rationale'],
                'rows': [[item.get('item', ''),
                         item.get('assessment', ''),
                         item.get('rationale', '')] for item in section['checklist_items']]
            }
            create_proper_table(doc, table_data)

    # Save document
    doc.save(output_path)
    return output_path

def detect_and_convert_tables(content):
    """Detect table-like content and convert to structured data"""

    tables_found = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Check for dash separators indicating a table
        if '-----' in line or '────' in line or '═══' in line:
            if i > 0:
                header_line = lines[i-1].strip()
                if header_line and not header_line.startswith('#'):
                    table_data = extract_table_from_position(lines, i-1)
                    if table_data:
                        tables_found.append(table_data)
                        content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")

        i += 1

    return content, tables_found

def extract_table_from_position(lines, start_idx):
    """Extract a table starting from a given position"""

    headers = []
    rows = []
    original_text_lines = []

    # Get headers
    header_line = lines[start_idx].strip()
    original_text_lines.append(lines[start_idx])

    # Split headers
    headers = re.split(r'\s{2,}|\t', header_line)
    headers = [h.strip() for h in headers if h.strip()]

    if len(headers) < 2:
        return None

    # Skip separator line
    if start_idx + 1 < len(lines):
        original_text_lines.append(lines[start_idx + 1])

    # Get data rows
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('#') or '═══' in line:
            break

        # Split row data
        row_data = re.split(r'\s{2,}|\t', line)
        row_data = [d.strip() for d in row_data if d.strip()]

        if row_data:
            while len(row_data) < len(headers):
                row_data.append('')
            rows.append(row_data[:len(headers)])
            original_text_lines.append(lines[i])

        i += 1

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def create_proper_table(doc, table_data):
    """Create a real Word table with borders"""

    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])

    if not headers:
        return

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # Set column widths
    try:
        available_width = Inches(6.5)
        col_width = available_width / len(headers)
        for col in table.columns:
            col.width = col_width
    except:
        pass

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
        # Gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D0D0D0')
        cell._element.get_or_add_tcPr().append(shading)

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data[:len(headers)]):
            cell = row.cells[i]
            cell.text = str(value) if value else ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

def main():
    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': 'Usage: script.py input.json output.docx'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        create_document_with_proper_numbering(data, output_path)
        print(json.dumps({'success': True}))

    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()