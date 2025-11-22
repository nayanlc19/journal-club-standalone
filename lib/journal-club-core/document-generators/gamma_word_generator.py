"""
Gamma-Ready Word Document Generator
Creates clean, presentation-optimized Word documents with text + images
Ready to upload directly to Gamma for auto-slide generation
"""

import sys
import json
import base64
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the lib directory to path to import universal sanitizer
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

try:
    # Try to import the universal sanitizer
    from universal_text_sanitizer import sanitize_text, detect_publisher
except ImportError:
    # Fallback to basic sanitizer if universal is not available
    def sanitize_text(text, context='general'):
        """Fallback sanitizer if universal is not available"""
        if not text:
            return text

        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

        # Replace special Unicode spaces with regular spaces
        text = text.replace('\u202f', ' ')  # Narrow no-break space
        text = text.replace('\u200b', '')   # Zero-width space
        text = text.replace('\u00a0', ' ')  # Non-breaking space
        text = text.replace('\ufeff', '')   # Zero-width no-break space

        return text

    def detect_publisher(text):
        return 'unknown'

def create_gamma_document(data):
    """
    Create a Gamma-ready Word document

    Args:
        data: {
            "title": "Study Title",
            "sections": [
                {
                    "heading": "Section Title",
                    "content": "Text content",
                    "images": [
                        {
                            "title": "Figure 1",
                            "base64": "...",
                            "explanation": "Brief explanation"
                        }
                    ]
                }
            ]
        }
    """
    doc = Document()

    # Title (sanitized)
    title = doc.add_heading(sanitize_text(data['title']), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process each section
    for section in data.get('sections', []):
        # Section heading (sanitized)
        doc.add_heading(sanitize_text(section['heading']), 1)

        # Section content (keep markdown formatting for readability)
        if section.get('content'):
            content = sanitize_text(section['content'])
            # Split content into paragraphs to preserve structure
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Check if this is a markdown table
                if '|' in para_text and para_text.count('\n') > 1:
                    # Keep table as preformatted text for Gamma AI to parse
                    table_para = doc.add_paragraph(para_text)
                    table_para.style = 'Normal'
                    for run in table_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                # Check if it's a heading (starts with ###)
                elif para_text.startswith('###'):
                    heading_text = para_text.replace('###', '').strip()
                    doc.add_heading(heading_text, 3)
                # Check if it's a bullet list
                elif para_text.startswith('- ') or para_text.startswith('* '):
                    lines = para_text.split('\n')
                    for line in lines:
                        if line.strip():
                            bullet_text = line.lstrip('- *').strip()
                            doc.add_paragraph(bullet_text, style='List Bullet')
                # Check if it's a numbered list
                elif len(para_text) > 2 and para_text[0].isdigit() and para_text[1] == '.':
                    lines = para_text.split('\n')
                    for line in lines:
                        if line.strip():
                            # Remove the number, let Word handle numbering
                            num_match = re.match(r'^\d+\.\s+(.+)', line)
                            if num_match:
                                doc.add_paragraph(num_match.group(1), style='List Number')
                            else:
                                doc.add_paragraph(line.strip())
                else:
                    # Regular paragraph
                    para = doc.add_paragraph(para_text)
                    para.style = 'Normal'

        # Add images if present
        for image_data in section.get('images', []):
            # Image title (sanitized)
            img_title = doc.add_paragraph(sanitize_text(image_data.get('title', 'Figure')))
            img_title.runs[0].bold = True
            img_title.runs[0].font.size = Pt(12)

            # Decode and insert image if available
            if image_data.get('base64'):
                try:
                    img_bytes = base64.b64decode(image_data['base64'])
                    img_stream = BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[Refer to original paper for this visual element]")
            else:
                # No image data - just note the reference
                doc.add_paragraph(f"[See original paper - page {image_data.get('pageNumber', '?')}]")

            # Image explanation (sanitized)
            if image_data.get('explanation'):
                explain = doc.add_paragraph(sanitize_text(image_data['explanation']))
                explain.runs[0].italic = True
                explain.runs[0].font.size = Pt(10)
                explain.runs[0].font.color.rgb = RGBColor(80, 80, 80)

            # Add spacing
            doc.add_paragraph()

    return doc


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python gamma_word_generator.py <input_json> <output_docx>"
        }))
        sys.exit(1)

    input_json = sys.argv[1]
    output_path = sys.argv[2]

    try:
        # Load input data
        with open(input_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Generate document
        doc = create_gamma_document(data)

        # Save to file
        doc.save(output_path)

        print(json.dumps({
            "success": True,
            "output": output_path,
            "message": "Gamma-ready Word document created successfully"
        }))

    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": str(e)
        }))
        sys.exit(1)


if __name__ == '__main__':
    main()
