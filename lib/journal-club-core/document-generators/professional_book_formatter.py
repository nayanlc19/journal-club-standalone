#!/usr/bin/env python3
"""
Professional Book-Style Document Formatter
Creates properly formatted Word documents with:
- Professional tables with borders and proper alignment
- Consistent paragraph spacing
- Book-style typography
- Clear section headers
- Proper indentation
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os
import re

class ProfessionalBookFormatter:
    """Creates professionally formatted Word documents with book-style layout"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.setup_page_layout()

    def setup_page_layout(self):
        """Configure page margins and layout for book-style formatting"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def setup_styles(self):
        """Define custom styles for professional formatting"""
        styles = self.doc.styles

        # Chapter title style
        chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Georgia'
        chapter_style.font.size = Pt(24)
        chapter_style.font.bold = True
        chapter_style.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        chapter_style.paragraph_format.space_before = Pt(0)
        chapter_style.paragraph_format.space_after = Pt(24)
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Section header style
        section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Georgia'
        section_style.font.size = Pt(16)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(18)
        section_style.paragraph_format.space_after = Pt(12)

        # Body text style
        body_style = styles.add_style('BookBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Garamond'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(8)
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.first_line_indent = Inches(0.3)

        # Result box style
        result_style = styles.add_style('ResultBox', WD_STYLE_TYPE.PARAGRAPH)
        result_style.font.name = 'Arial'
        result_style.font.size = Pt(10)
        result_style.font.italic = True
        result_style.paragraph_format.left_indent = Inches(0.5)
        result_style.paragraph_format.right_indent = Inches(0.5)
        result_style.paragraph_format.space_before = Pt(6)
        result_style.paragraph_format.space_after = Pt(6)

        # Interpretation style
        interp_style = styles.add_style('Interpretation', WD_STYLE_TYPE.PARAGRAPH)
        interp_style.font.name = 'Calibri'
        interp_style.font.size = Pt(11)
        interp_style.paragraph_format.space_before = Pt(6)
        interp_style.paragraph_format.space_after = Pt(12)
        interp_style.paragraph_format.first_line_indent = Inches(0.3)

    def add_chapter_title(self, title):
        """Add a chapter title with proper formatting"""
        p = self.doc.add_paragraph(title, style='ChapterTitle')
        return p

    def add_section_header(self, header):
        """Add a section header with proper formatting"""
        p = self.doc.add_paragraph(header, style='SectionHeader')
        return p

    def add_body_text(self, text):
        """Add body text with book-style formatting"""
        # Clean up text
        text = text.strip()
        if text:
            p = self.doc.add_paragraph(text, style='BookBody')
            return p

    def add_result_box(self, result_text):
        """Add a highlighted result box"""
        # Add a shaded paragraph for results
        p = self.doc.add_paragraph(style='ResultBox')

        # Add shading
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E8E8E8')  # Light gray background
        p._element.get_or_add_pPr().append(shading_elm)

        # Add "Result:" in bold
        run = p.add_run('Result: ')
        run.bold = True

        # Add the actual result text
        p.add_run(result_text)

        return p

    def add_interpretation(self, text):
        """Add interpretation text with proper formatting"""
        p = self.doc.add_paragraph(style='Interpretation')

        # Add "Interpretation:" in bold italic
        run = p.add_run('Interpretation: ')
        run.bold = True
        run.italic = True

        # Add the interpretation text
        p.add_run(text)

        return p

    def create_professional_table(self, headers, data, caption=None):
        """Create a professionally formatted table with borders"""

        if caption:
            # Add table caption
            caption_p = self.doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_p.add_run(f'Table: {caption}')
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_p.paragraph_format.space_after = Pt(6)

        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'  # This adds borders
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths (distribute evenly)
        table.autofit = False
        col_width = Inches(6.0 / len(headers))  # Total width divided by columns
        for col in table.columns:
            col.width = col_width

        # Add headers with formatting
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header)

            # Format header cell
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Bold header text
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)

            # Add gray background to header
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')  # Light gray
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data[:len(headers)]):
                cell = row_cells[i]
                cell.text = str(value) if value else ''

                # Format data cells
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Set font size
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        return table

    def format_statistical_results(self, stats_text):
        """Format statistical results with proper notation"""
        # Clean up statistical notation
        formatted = stats_text

        # Format p-values
        formatted = re.sub(r'p\s*([<>=])\s*([\d.]+)', r'p \1 \2', formatted)

        # Format confidence intervals
        formatted = re.sub(r'CI\s*:\s*([\d.]+)\s*-\s*([\d.]+)', r'CI: \1–\2', formatted)

        # Format percentages
        formatted = re.sub(r'(\d+)\s*%', r'\1%', formatted)

        return formatted

    def add_key_findings_box(self, findings):
        """Add a highlighted box for key findings"""
        # Add "Key Findings" header
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('KEY FINDINGS')
        run.font.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        # Create a table for the box effect
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)

        # Add findings as bullet points
        for finding in findings:
            p = cell.add_paragraph(f'• {finding}')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)

        # Shade the cell
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F0F0')  # Very light gray
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Add spacing after
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    def save(self, filename):
        """Save the document with professional formatting"""
        self.doc.save(filename)
        print(f"Professional document saved: {filename}")

def format_appraisal_to_book(appraisal_data, output_path):
    """Convert appraisal data to professional book-format document"""

    formatter = ProfessionalBookFormatter()

    # Add title
    formatter.add_chapter_title("CRITICAL APPRAISAL")

    # Add subtitle with paper details
    if 'title' in appraisal_data:
        p = formatter.doc.add_paragraph(appraisal_data['title'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(12)

    # Add PICO question
    if 'clinical_question' in appraisal_data:
        formatter.add_section_header("Clinical Question (PICO)")
        formatter.add_body_text(appraisal_data['clinical_question'])

    # Add study design
    if 'study_type' in appraisal_data:
        formatter.add_section_header("Study Design")
        formatter.add_body_text(f"Type: {appraisal_data['study_type']}")

    # Add results section
    formatter.add_chapter_title("CHAPTER 8")
    formatter.add_section_header("RESULTS INTERPRETATION (SIMPLE LANGUAGE)")

    # Add primary outcome
    if 'primary_outcome' in appraisal_data:
        formatter.add_body_text("Primary outcome:")
        formatter.add_body_text(f"Overall Caesarean-section rate")

        formatter.add_result_box("165 CS out of 352 total deliveries → 46.9% (rounded to one decimal).")

        formatter.add_interpretation(
            "Almost one in two women who gave birth in this hospital during the study period had a caesarean "
            "section. In other words, for every 100 deliveries, about 47 were performed by surgery rather than "
            "vaginally."
        )

    # Add any tables with proper formatting
    if 'tables' in appraisal_data:
        for table_data in appraisal_data['tables']:
            if 'headers' in table_data and 'rows' in table_data:
                formatter.create_professional_table(
                    headers=table_data['headers'],
                    data=table_data['rows'],
                    caption=table_data.get('caption', '')
                )

    # Add key findings box
    if 'key_findings' in appraisal_data:
        formatter.add_key_findings_box(appraisal_data['key_findings'])

    # Add statistical analysis section
    if 'statistical_test' in appraisal_data:
        formatter.add_section_header("Statistical Analysis")
        formatted_stats = formatter.format_statistical_results(appraisal_data['statistical_test'])
        formatter.add_body_text(formatted_stats)

    # Save the document
    formatter.save(output_path)

    return output_path

# Example usage
if __name__ == "__main__":
    # Test data
    test_data = {
        'title': 'Effect of Robson Classification on Caesarean Section Rates',
        'clinical_question': 'In women delivering at tertiary care hospitals, does the Robson classification system reduce caesarean section rates?',
        'study_type': 'Prospective cohort study',
        'primary_outcome': 'Caesarean section rate',
        'tables': [
            {
                'caption': 'Distribution of CS by Robson Groups',
                'headers': ['Robson Group', 'Number of CS', '% of all CS', '% of all deliveries'],
                'rows': [
                    ['1', '10', '6.1%', '2.8%'],
                    ['2', '33', '20.0%', '9.4%'],
                    ['3', '7', '4.2%', '2.0%'],
                    ['4', '16', '9.7%', '4.5%'],
                    ['5', '41', '24.8%', '11.6%']
                ]
            }
        ],
        'key_findings': [
            'CS rate was 46.9% overall',
            'Group 5 (previous CS) contributed most to CS rate',
            'Significant variation between Robson groups (p < 0.001)'
        ],
        'statistical_test': 'Chi-square test: p < 0.001'
    }

    output_file = 'professional_appraisal.docx'
    format_appraisal_to_book(test_data, output_file)