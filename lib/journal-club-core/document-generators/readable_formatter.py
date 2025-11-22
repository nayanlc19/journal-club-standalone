#!/usr/bin/env python3
"""
Readable formatter with clear visual separation and proper structure
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import sys
import re

def add_shaded_box(doc, content, bg_color='F0F0F0'):
    """Add a shaded box for better visual separation"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    # Add content to cell
    p = cell.paragraphs[0]
    p.text = content
    p.paragraph_format.space_after = Pt(3)

    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    cell._element.get_or_add_tcPr().append(shading)

    # Add spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def format_checklist_item(doc, item_data):
    """Format a checklist item with clear visual separation"""

    # Create a table for the checklist item (better visual structure)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    # Add the check symbol and question
    if 'item' in item_data:
        # Question/Item header
        p = cell.add_paragraph()

        # Add symbol (✓, ⚠, ❌)
        symbol = '✓' if item_data.get('assessment') == 'YES' else '⚠' if item_data.get('assessment') == 'CANT TELL' else '❌'
        run = p.add_run(f"{symbol} ")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 128, 0) if symbol == '✓' else RGBColor(255, 165, 0) if symbol == '⚠' else RGBColor(255, 0, 0)

        # Add the question
        run = p.add_run(item_data['item'])
        run.font.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    # Add Answer in a distinct format
    if 'assessment' in item_data:
        p = cell.add_paragraph()
        run = p.add_run("Answer: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        run = p.add_run(item_data['assessment'])
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Add Evidence in a distinct format
    if 'rationale' in item_data and item_data['rationale']:
        p = cell.add_paragraph()
        run = p.add_run("Evidence: ")
        run.font.bold = True
        run.font.italic = True
        run = p.add_run(item_data['rationale'])
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Add "Why it matters" if present
    if 'why_it_matters' in item_data:
        p = cell.add_paragraph()
        run = p.add_run("Why it matters: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 0, 128)
        run = p.add_run(item_data['why_it_matters'])
        run.font.size = Pt(10)
        run.font.italic = True
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Add background color for better separation
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F8F8F8')
    cell._element.get_or_add_tcPr().append(shading)

    # Add spacing after the item
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

def format_methods_section(doc, content):
    """Format methods section with clear Q&A structure"""

    lines = content.split('\n')
    current_section = []

    for line in lines:
        line = line.strip()

        # Check if this is a numbered heading
        if re.match(r'^\d+\.\s+.+', line):
            # Process previous section if exists
            if current_section:
                process_qa_section(doc, current_section)
                current_section = []

            # Add the numbered heading with proper formatting
            p = doc.add_paragraph()
            p.add_run(line)
            p.style = 'Heading 3'
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)

            # Add a horizontal line for separation
            add_separator_line(doc)

        else:
            current_section.append(line)

    # Process last section
    if current_section:
        process_qa_section(doc, current_section)

def process_qa_section(doc, lines):
    """Process a Q&A section with proper formatting"""

    text = '\n'.join(lines)

    # Look for Q&A patterns
    if 'What' in text or 'Why' in text or 'How' in text:
        # This looks like a question
        parts = text.split(':', 1)
        if len(parts) == 2:
            # Format as Q&A
            # Question in a box
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            run = p.add_run(parts[0] + ':')
            run.font.bold = True
            run.font.size = Pt(11)

            # Answer below with indentation
            p = cell.add_paragraph()
            p.add_run(parts[1].strip())
            p.paragraph_format.left_indent = Inches(0.3)

            # Light blue background for Q&A
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E6F3FF')
            cell._element.get_or_add_tcPr().append(shading)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        else:
            # Regular paragraph
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
    else:
        # Regular text - check for bullet points
        for line in lines:
            if line.strip():
                if line.strip().startswith('•') or line.strip().startswith('-'):
                    # Bullet point
                    p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(4)
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Inches(0.25)

def add_separator_line(doc):
    """Add a visual separator line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    # Add bottom border to create a line effect
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._element.get_or_add_pPr().append(pBdr)

def create_readable_document(data, output_path):
    """Create a document with readable formatting"""

    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = data.get('title', 'Critical Appraisal')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(18)

    # Process sections
    for section in data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')

        # Add section heading with visual separation
        if heading:
            # Add a separator before major sections
            if 'Critical Appraisal Checklist' in heading or 'Methods Deep-Dive' in heading:
                add_separator_line(doc)

            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)

        # Process content based on section type
        if 'checklist_items' in section and section['checklist_items']:
            # Format checklist with clear separation
            for item in section['checklist_items']:
                format_checklist_item(doc, item)

        elif 'Methods' in heading and content:
            # Format methods section with Q&A structure
            format_methods_section(doc, content)

        elif content:
            # Process regular content
            # Detect and convert tables
            cleaned_content, detected_tables = detect_and_convert_tables(content)

            # Process paragraphs
            paragraphs = cleaned_content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Check for table markers
                if para.startswith('[TABLE_') and para.endswith(']'):
                    try:
                        table_num = int(para[7:-1])
                        if table_num < len(detected_tables):
                            create_proper_table(doc, detected_tables[table_num])
                    except:
                        pass

                # Check for Result/Interpretation blocks
                elif 'Result:' in para or 'Interpretation:' in para:
                    # Create a visually distinct block
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)

                    if 'Result:' in para:
                        parts = para.split('Result:', 1)
                        if len(parts) > 1:
                            p = cell.paragraphs[0]
                            run = p.add_run('Result:')
                            run.font.bold = True
                            run.font.size = Pt(11)
                            p.add_run('\n' + parts[1].strip())

                            # Light gray background
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'E8E8E8')
                            cell._element.get_or_add_tcPr().append(shading)

                    if 'Interpretation:' in para:
                        parts = para.split('Interpretation:', 1)
                        if len(parts) > 1:
                            p = cell.add_paragraph() if 'Result:' in para else cell.paragraphs[0]
                            run = p.add_run('Interpretation:')
                            run.font.bold = True
                            run.font.italic = True
                            run.font.size = Pt(11)
                            p.add_run('\n' + parts[1].strip())

                    doc.add_paragraph().paragraph_format.space_after = Pt(8)

                else:
                    # Regular paragraph with proper spacing
                    p = doc.add_paragraph()
                    p.add_run(para)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(8)

                    # Add indentation for body text
                    if not para.startswith('•') and not para.startswith('-'):
                        p.paragraph_format.first_line_indent = Inches(0.25)

    # Save document
    doc.save(output_path)
    return output_path

def detect_and_convert_tables(content):
    """Detect table-like content and convert to structured data"""

    tables_found = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Check for dash separators
        if '-----' in line or '────' in line:
            if i > 0:
                header_line = lines[i-1].strip()
                if header_line and not header_line.startswith('#'):
                    table_data = extract_table_from_position(lines, i-1)
                    if table_data:
                        tables_found.append(table_data)
                        content = content.replace(table_data['original_text'], f"[TABLE_{len(tables_found)-1}]")

        i += 1

    return content, tables_found

def extract_table_from_position(lines, start_idx):
    """Extract a table starting from a given position"""

    headers = []
    rows = []
    original_text_lines = []

    # Get headers
    header_line = lines[start_idx].strip()
    original_text_lines.append(lines[start_idx])

    # Split headers
    headers = re.split(r'\s{2,}|\t', header_line)
    headers = [h.strip() for h in headers if h.strip()]

    if len(headers) < 2:
        return None

    # Skip separator line
    if start_idx + 1 < len(lines):
        original_text_lines.append(lines[start_idx + 1])

    # Get data rows
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('#'):
            break

        row_data = re.split(r'\s{2,}|\t', line)
        row_data = [d.strip() for d in row_data if d.strip()]

        if row_data:
            while len(row_data) < len(headers):
                row_data.append('')
            rows.append(row_data[:len(headers)])
            original_text_lines.append(lines[i])

        i += 1

    if rows:
        return {
            'headers': headers,
            'rows': rows,
            'original_text': '\n'.join(original_text_lines)
        }

    return None

def create_proper_table(doc, table_data):
    """Create a real Word table with borders"""

    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])

    if not headers:
        return

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # Gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D0D0D0')
        cell._element.get_or_add_tcPr().append(shading)

    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data[:len(headers)]):
            cell = row.cells[i]
            cell.text = str(value) if value else ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

def main():
    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': 'Usage: script.py input.json output.docx'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        create_readable_document(data, output_path)
        print(json.dumps({'success': True}))

    except Exception as e:
        print(json.dumps({'success': False, 'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()